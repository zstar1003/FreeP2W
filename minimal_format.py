"""
Minimal formatting - Only fix critical issues, preserve original formatting
Based on actual smallpdf behavior: preserve fonts, spacing, indents
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def minimal_format_docx(input_path, output_path):
    """
    Minimal formatting:
    1. Unify body text font size to 10pt (only if inconsistent)
    2. Clean TOC spacing (remove extra spaces)
    3. Preserve original fonts, alignment, spacing, indents
    4. Don't change existing alignment
    """

    doc = Document(input_path)

    print("=" * 80)
    print("MINIMAL FORMATTING (Preserving Original Format)")
    print("=" * 80)

    # Detect TOC
    toc_start = None
    toc_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if toc_start is None and ('contents' in text or 'table of contents' in text):
            toc_start = i
            print(f"\nTOC starts at paragraph {i}")

        if toc_start is not None and toc_end is None and i > toc_start:
            para_text = para.text.strip()
            if (re.match(r'^\d+\.?\d*\s+[A-Z]', para_text) or
                any(keyword in text for keyword in ['abstract', 'introduction']) or
                len(para_text) > 100):
                toc_end = i
                print(f"TOC ends at paragraph {i}")
                break

    if toc_start and not toc_end:
        toc_end = min(toc_start + 15, len(doc.paragraphs))

    # Collect font sizes to determine body text size
    font_sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size and run.font.size.pt <= 12:
                font_sizes.append(run.font.size.pt)

    # Determine target body size (most common size <= 12pt)
    from collections import Counter
    if font_sizes:
        target_body_size = Counter(font_sizes).most_common(1)[0][0]
    else:
        target_body_size = 10.0

    print(f"Target body text size: {target_body_size}pt")

    # Minimal formatting
    stats = {
        'toc_cleaned': 0,
        'size_unified': 0,
        'preserved': 0,
        'math_skipped': 0
    }

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Skip math paragraphs
        has_math = 'oMath' in para._element.xml if hasattr(para, '_element') else False
        if has_math:
            stats['math_skipped'] += 1
            continue

        # Only clean TOC spacing
        in_toc = toc_start is not None and toc_start <= i < toc_end

        if in_toc:
            # Clean extra spaces in TOC
            cleaned = re.sub(r'\s+', ' ', text).strip()
            if cleaned != text and para.runs:
                # Preserve original font, just clean text
                original_font = para.runs[0].font.name if para.runs[0].font.name else None
                original_size = para.runs[0].font.size if para.runs[0].font.size else None

                para.clear()
                run = para.add_run(cleaned)
                if original_font:
                    run.font.name = original_font
                if original_size:
                    run.font.size = original_size
                else:
                    run.font.size = Pt(target_body_size)

                stats['toc_cleaned'] += 1
            else:
                stats['preserved'] += 1

        else:
            # For body text: only unify font size if it's close to target (9-11pt range)
            # Preserve font names, alignment, spacing, indents!
            size_changed = False

            for run in para.runs:
                if run.font.size:
                    size_pt = run.font.size.pt
                    # Only unify if size is in the "body text" range and not exact
                    if 9 <= size_pt <= 11 and size_pt != target_body_size:
                        run.font.size = Pt(target_body_size)
                        size_changed = True

            if size_changed:
                stats['size_unified'] += 1
            else:
                stats['preserved'] += 1

    print(f"\nFormatting summary:")
    print(f"  - TOC entries cleaned: {stats['toc_cleaned']}")
    print(f"  - Font sizes unified: {stats['size_unified']}")
    print(f"  - Paragraphs preserved: {stats['preserved']}")
    print(f"  - Math formulas skipped: {stats['math_skipped']}")

    # Save
    doc.save(output_path)
    print(f"\nSaved to: {output_path}")
    print("=" * 80)

if __name__ == "__main__":
    minimal_format_docx("result.docx", "result_minimal_format.docx")
