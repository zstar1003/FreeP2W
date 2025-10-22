"""
Check if the text exists in pure pdf2docx output (without hybrid method)
"""
from docx import Document

docx_path = "test_files/result_pdf2docx_origin.docx"

doc = Document(docx_path)

with open("check_original_result.txt", "w", encoding="utf-8") as f:
    f.write("Searching for 'groud truth velocity' in pure pdf2docx output...\n\n")
    f.write("=" * 80 + "\n")

    found_count = 0

    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        if 'groud truth' in text.lower() or ('velocity' in text.lower() and 'vt is' in text.lower()):
            found_count += 1
            f.write(f"\nParagraph {idx}:\n")
            f.write(f"  Text: {text[:300]}\n")

    # Also check context around formula paragraphs
    f.write("\n\nParagraphs around formulas (35-40):\n")
    f.write("=" * 80 + "\n")

    for idx in range(30, min(45, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        if text:
            f.write(f"\nPara {idx}: {text[:150]}\n")

    f.write("\n" + "=" * 80 + "\n")
    f.write(f"Total 'groud truth velocity' occurrences: {found_count}\n")

    if found_count == 0:
        f.write("\n[FINDING] Text 'The groud truth velocity vt is' is ALSO MISSING in pure pdf2docx output!\n")
        f.write("This suggests the problem is with pdf2docx, NOT our hybrid method.\n")
    else:
        f.write("\n[FINDING] Text exists in pure pdf2docx but missing in hybrid output!\n")
        f.write("This suggests our hybrid method is causing the problem.\n")

    f.write("=" * 80 + "\n")

print("Results written to check_original_result.txt")
