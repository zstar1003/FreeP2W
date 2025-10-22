"""
Verify table-based formula layout
"""
from docx import Document

docx_path = "test_files/result_hybrid.docx"

doc = Document(docx_path)

print("Checking for tables in the document...\n")
print("=" * 80)

table_count = 0
for idx, table in enumerate(doc.tables):
    # Check if this is a formula table (1 row, 3 columns)
    if len(table.rows) == 1 and len(table.columns) == 3:
        table_count += 1
        row = table.rows[0]

        print(f"\nTable {idx} (Formula Table #{table_count}):")
        print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")

        # Cell 0: Empty
        cell0_text = row.cells[0].text.strip()
        print(f"  Cell 0 (Empty): '{cell0_text}'")

        # Cell 1: Formula (should have m:oMath)
        cell1 = row.cells[1]
        cell1_xml = cell1._element.xml.decode('utf-8') if isinstance(cell1._element.xml, bytes) else str(cell1._element.xml)
        has_math = 'm:oMath' in cell1_xml
        print(f"  Cell 1 (Formula): Has math = {has_math}")
        if cell1.paragraphs:
            print(f"    Alignment: {cell1.paragraphs[0].alignment}")

        # Cell 2: Equation number
        cell2 = row.cells[2]
        cell2_text = cell2.text.strip()
        print(f"  Cell 2 (Number): '{cell2_text}'")
        if cell2.paragraphs:
            print(f"    Alignment: {cell2.paragraphs[0].alignment}")

print(f"\n" + "=" * 80)
print(f"Total formula tables found: {table_count}")
print("=" * 80)
