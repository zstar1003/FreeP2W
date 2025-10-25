"""
Hybrid PDF to DOCX Converter
Combines DocLayout-YOLO (for images and formulas) + pdf2docx (for text/tables)

Strategy:
1. Use DocLayout-YOLO to detect figure and formula regions
2. Use UniMERNet to recognize formulas and convert to MathML
3. Use pdf2docx for text, tables, and other content
4. Replace pdf2docx's image extraction with YOLO-detected regions
"""

import os
import torch
from doclayout_yolo import YOLOv10
from pdf2docx import Converter
import fitz


class YOLORegionDetector:
    """DocLayout-YOLO region detector for pdf2docx integration (figures + formulas)"""

    def __init__(self, model_path="weights/doclayout_yolo_docstructbench_imgsz1024.pt",
                 imgsz=1024, conf=0.25, dpi=150):
        """
        Initialize YOLO detector

        Args:
            model_path: Path to YOLO model
            imgsz: Image size for detection
            conf: Confidence threshold
            dpi: DPI for PDF rendering
        """
        print(f"[YOLO] Loading DocLayout-YOLO model...")

        # Auto select device
        self.device = 'cuda' if torch.cuda.is_available() else 'cpu'
        print(f"[YOLO] Using device: {self.device}")

        # Load model
        self.model = YOLOv10(model_path)
        self.imgsz = imgsz
        self.conf = conf
        self.dpi = dpi

        # Cache for detected regions (page_num -> {'figures': [...], 'formulas': [...]})
        self.detected_regions = {}

        print(f"[YOLO] Model loaded successfully\n")

    def detect_page_regions(self, page, page_num):
        """
        Detect figure and formula regions in a PDF page

        Args:
            page: fitz.Page object
            page_num: Page number (0-indexed)

        Returns:
            dict: {'figures': [regions], 'formulas': [regions]}
                  Each region: {'bbox': [x0,y0,x1,y1], 'confidence': float}
        """
        # Check cache
        if page_num in self.detected_regions:
            return self.detected_regions[page_num]

        # Convert page to image
        mat = fitz.Matrix(self.dpi / 72, self.dpi / 72)
        pix = page.get_pixmap(matrix=mat)

        # Save temp image
        temp_img = f"_temp_yolo_page_{page_num}.png"
        pix.save(temp_img)

        # Run YOLO detection
        results = self.model.predict(
            temp_img,
            imgsz=self.imgsz,
            conf=self.conf,
            device=self.device,
            verbose=False  # Suppress YOLO output
        )

        # Clean up temp file
        if os.path.exists(temp_img):
            os.remove(temp_img)

        # Parse results
        figures = []
        formulas = []

        if len(results) > 0 and hasattr(results[0], 'boxes'):
            for box in results[0].boxes:
                # Get class
                cls_id = int(box.cls[0].cpu().numpy())
                class_name = self.model.names[cls_id]

                # Get bbox in image coordinates
                x1, y1, x2, y2 = box.xyxy[0].cpu().numpy()

                # Convert to PDF coordinates
                scale_x = page.rect.width / pix.width
                scale_y = page.rect.height / pix.height

                pdf_bbox = [
                    float(x1 * scale_x),
                    float(y1 * scale_y),
                    float(x2 * scale_x),
                    float(y2 * scale_y)
                ]

                region_data = {
                    'bbox': pdf_bbox,
                    'confidence': float(box.conf[0].cpu().numpy())
                }

                # Categorize by type
                if class_name.lower() in ['figure', 'picture']:
                    figures.append(region_data)
                elif class_name.lower() in ['isolate_formula', 'formula']:
                    formulas.append(region_data)

        # Cache results
        regions = {'figures': figures, 'formulas': formulas}
        self.detected_regions[page_num] = regions

        if figures or formulas:
            print(f"[YOLO] Page {page_num + 1}: Detected {len(figures)} figures, {len(formulas)} formulas")

        return regions


class HybridConverter:
    """
    Hybrid PDF to DOCX Converter
    Uses DocLayout-YOLO for images and formulas, pdf2docx for everything else
    """

    def __init__(self, yolo_model_path="weights/doclayout_yolo_docstructbench_imgsz1024.pt",
                 unimernet_cfg_path="demo.yaml"):
        """
        Initialize hybrid converter

        Args:
            yolo_model_path: Path to DocLayout-YOLO model
            unimernet_cfg_path: Path to UniMERNet config file (default: demo.yaml in project root)
        """
        self.yolo_detector = YOLORegionDetector(model_path=yolo_model_path)
        self.unimernet_cfg_path = unimernet_cfg_path
        self.formula_recognizer = None  # Lazy load UniMERNet
        self.formula_processor = None

    def _load_unimernet(self):
        """Lazy load UniMERNet model"""
        if self.formula_recognizer is not None:
            return

        print("[UniMERNet] Loading formula recognition model...")

        import sys
        import os

        # Get project root directory (where hybrid_converter.py is located)
        project_root = os.path.dirname(os.path.abspath(__file__))

        # Add project root to path if not already there
        if project_root not in sys.path:
            sys.path.insert(0, project_root)

        try:
            import argparse
            import unimernet.tasks as tasks
            from unimernet.common.config import Config
            from unimernet.processors import load_processor

            # Setup model - use absolute path for config
            cfg_path = os.path.join(project_root, self.unimernet_cfg_path)
            if not os.path.exists(cfg_path):
                raise FileNotFoundError(f"Config file not found: {cfg_path}")

            args = argparse.Namespace(cfg_path=cfg_path, options=None)
            cfg = Config(args)

            task = tasks.setup_task(cfg)
            device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

            # Build and load model
            self.formula_recognizer = task.build_model(cfg).to(device)
            self.formula_recognizer.eval()

            # Load processor
            self.formula_processor = load_processor(
                "formula_image_eval",
                cfg.config.datasets.formula_rec_eval.vis_processor.eval,
            )

            print(f"[UniMERNet] Model loaded on {device}\n")

        except ImportError as e:
            print(f"[Error] Failed to import UniMERNet modules: {e}")
            print(f"[Error] Make sure 'unimernet' folder is in: {project_root}")
            raise
        except Exception as e:
            print(f"[Error] Failed to load UniMERNet model: {e}")
            raise

    def _recognize_formula(self, page, bbox):
        """
        Recognize formula from PDF page region

        Args:
            page: fitz.Page object
            bbox: [x0, y0, x1, y1] coordinates

        Returns:
            str: LaTeX formula string
        """
        # Ensure UniMERNet is loaded
        self._load_unimernet()

        # Extract formula region as image
        rect = fitz.Rect(bbox)
        mat = fitz.Matrix(3.0, 3.0)  # Higher resolution for formula recognition
        pix = page.get_pixmap(matrix=mat, clip=rect)

        # Convert to PIL Image
        import io
        from PIL import Image
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data)).convert("RGB")

        # Process with UniMERNet
        device = next(self.formula_recognizer.parameters()).device
        image_tensor = self.formula_processor(img).unsqueeze(0).to(device)

        with torch.no_grad():
            output = self.formula_recognizer.generate({"image": image_tensor})

        latex_code = output["pred_str"][0]
        return latex_code

    def _latex_to_mathml(self, latex_code):
        """
        Convert LaTeX to MathML

        Args:
            latex_code: LaTeX string

        Returns:
            str: MathML string
        """
        try:
            from latex2mathml.converter import convert
            mathml = convert(latex_code)
            return mathml
        except Exception as e:
            print(f"[Warning] LaTeX to MathML conversion failed: {e}")
            print(f"  LaTeX: {latex_code}")
            return None

    def convert(self, pdf_path, docx_path, **kwargs):
        """
        Convert PDF to DOCX with hybrid approach

        Args:
            pdf_path: Input PDF file
            docx_path: Output DOCX file
            **kwargs: Additional parameters for pdf2docx

        Supported kwargs:
            - use_yolo_images: bool (default True) - Use YOLO for image detection
            - yolo_conf: float (default 0.25) - YOLO confidence threshold
            - yolo_imgsz: int (default 1024) - YOLO image size
            - ... (other pdf2docx parameters)
        """
        use_yolo = kwargs.pop('use_yolo_images', True)

        # Update YOLO parameters if provided
        if 'yolo_conf' in kwargs:
            self.yolo_detector.conf = kwargs.pop('yolo_conf')
        if 'yolo_imgsz' in kwargs:
            self.yolo_detector.imgsz = kwargs.pop('yolo_imgsz')

        print(f"\n{'='*80}")
        print(f"Hybrid Conversion: {pdf_path}")
        print(f"  - Images & Formulas: {'DocLayout-YOLO + UniMERNet' if use_yolo else 'pdf2docx'}")
        print(f"  - Text/Tables: pdf2docx")
        print(f"{'='*80}\n")

        if use_yolo:
            # Pre-detect all figures and formulas with YOLO
            print("[YOLO] Pre-detecting figures and formulas in all pages...")
            pdf_doc = fitz.open(pdf_path)

            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                self.yolo_detector.detect_page_regions(page, page_num)

            pdf_doc.close()
            print()

            # Inject YOLO detector into pdf2docx
            self._inject_yolo_to_pdf2docx(pdf_path)

        # Convert with pdf2docx
        print("[pdf2docx] Converting with pdf2docx...")
        cv = Converter(pdf_path)

        # Set default parameters for better quality
        default_params = {
            'start': 0,
            'end': None,
        }
        default_params.update(kwargs)

        cv.convert(docx_path, **default_params)
        cv.close()

        # Insert formulas if any were detected
        if use_yolo:
            self._insert_formulas_to_docx(pdf_path, docx_path)

        print(f"\n{'='*80}")
        print(f"Conversion completed!")
        print(f"Output: {docx_path}")
        print(f"{'='*80}\n")

    def _insert_formulas_to_docx(self, pdf_path, docx_path):
        """
        Insert MathML formulas into the generated DOCX file

        Args:
            pdf_path: Original PDF file
            docx_path: Generated DOCX file
        """
        # Check if there are any formulas detected
        has_formulas = False
        for regions in self.yolo_detector.detected_regions.values():
            if regions.get('formulas'):
                has_formulas = True
                break

        if not has_formulas:
            return

        print("\n[Formula] Processing formulas...")

        # Open PDF
        pdf_doc = fitz.open(pdf_path)

        # Recognize all formulas
        formula_data = {}  # page_num -> [(bbox, latex, mathml), ...]

        for page_num, regions in self.yolo_detector.detected_regions.items():
            formulas = regions.get('formulas', [])
            if not formulas:
                continue

            page = pdf_doc[page_num]
            page_formulas = []

            for formula_idx, formula in enumerate(formulas):
                bbox = formula['bbox']

                # Recognize formula
                print(f"[Formula] Page {page_num + 1}: Recognizing formula at {bbox}...")
                latex = self._recognize_formula(page, bbox)
                print(f"  LaTeX: {latex}")

                # Convert to MathML
                mathml = self._latex_to_mathml(latex)

                if mathml:
                    formula_info = {
                        'bbox': bbox,
                        'latex': latex,
                        'mathml': mathml
                    }

                    # Preserve equation number if detected
                    if 'equation_number' in formula:
                        formula_info['equation_number'] = formula['equation_number']
                        print(f"  Equation number: {formula['equation_number']}")

                    page_formulas.append(formula_info)
                    print(f"  ✓ Converted to MathML")
                else:
                    print(f"  ✗ MathML conversion failed")

            if page_formulas:
                formula_data[page_num] = page_formulas

        pdf_doc.close()

        # Insert formulas into DOCX
        if formula_data:
            print(f"\n[Formula] Inserting {sum(len(f) for f in formula_data.values())} formulas into DOCX...")
            self._replace_formula_placeholders(docx_path, formula_data)
            print(f"[Formula] Formulas inserted successfully\n")

    def _replace_formula_placeholders(self, docx_path, formula_data):
        """
        Replace formula placeholder text with actual MathML formulas in DOCX file

        Args:
            docx_path: Generated DOCX file
            formula_data: Dict of {page_num: [{'bbox': [...], 'latex': str, 'mathml': str}]}
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        # Open document
        doc = Document(docx_path)

        # Create a mapping of placeholder IDs to formula data
        formula_map = {}
        for page_num, formulas in formula_data.items():
            for formula_idx, formula_info in enumerate(formulas):
                placeholder_id = f"FORMULA_PLACEHOLDER_{page_num}_{formula_idx}"
                formula_map[placeholder_id] = formula_info

        print(f"[Formula] Searching for {len(formula_map)} placeholder texts in document...")

        replaced_count = 0

        # Iterate through all paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text

            # Look for formula placeholder pattern
            for placeholder_id, formula_info in formula_map.items():
                if placeholder_id in para_text:
                    mathml = formula_info['mathml']
                    latex = formula_info['latex']
                    equation_number = formula_info.get('equation_number', None)

                    print(f"  [Formula] Found placeholder '{placeholder_id}' at paragraph {para_idx}, replacing...")

                    try:
                        # Extract text before and after the placeholder
                        placeholder_pos = para_text.find(placeholder_id)
                        text_before = para_text[:placeholder_pos].strip()
                        text_after = para_text[placeholder_pos + len(placeholder_id):].strip()

                        # Remove equation number from text_after if present
                        # (equation number is already in the table, no need to duplicate)
                        if equation_number and text_after:
                            import re
                            # Remove all occurrences of equation number pattern
                            pattern = re.escape(equation_number)
                            # Remove pattern with surrounding whitespace/tabs (global replace)
                            text_after = re.sub(r'[\s\t]*' + pattern + r'[\s\t]*', '', text_after)
                            # Clean up any remaining tabs or excessive whitespace
                            text_after = re.sub(r'\t+', ' ', text_after).strip()
                            # If nothing left, set to empty
                            if not text_after or text_after.isspace():
                                text_after = ''

                        # Strip MathML declaration if present
                        if mathml.startswith('<?xml'):
                            mathml = mathml[mathml.index('?>') + 2:].strip()

                        # Create OMML element from MathML
                        math_run = self._create_omml_from_mathml(mathml)

                        if math_run:
                            # Use a single-row, 3-column invisible table for perfect alignment
                            # This is the standard professional method in Word for centered formula + right-aligned number
                            from docx.shared import Pt, Inches
                            from docx.oxml import OxmlElement
                            from docx.oxml.ns import qn
                            from docx.enum.text import WD_ALIGN_PARAGRAPH

                            # Create a table structure in XML (1 row, 3 columns)
                            # Column 1: Empty (balance)
                            # Column 2: Formula (centered)
                            # Column 3: Number (right-aligned)

                            tbl = OxmlElement('w:tbl')

                            # Table properties - no borders, full text width
                            tblPr = OxmlElement('w:tblPr')

                            # Table width = 100% of text width (between margins)
                            tblW = OxmlElement('w:tblW')
                            tblW.set(qn('w:type'), 'pct')
                            tblW.set(qn('w:w'), '5000')  # 100% = 5000 in pct
                            tblPr.append(tblW)

                            # Table indent = 0 (align with text left margin)
                            tblInd = OxmlElement('w:tblInd')
                            tblInd.set(qn('w:w'), '0')
                            tblInd.set(qn('w:type'), 'dxa')
                            tblPr.append(tblInd)

                            # No borders
                            tblBorders = OxmlElement('w:tblBorders')
                            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'none')
                                border.set(qn('w:sz'), '0')
                                border.set(qn('w:space'), '0')
                                border.set(qn('w:color'), 'auto')
                                tblBorders.append(border)
                            tblPr.append(tblBorders)

                            # Table layout: fixed
                            tblLayout = OxmlElement('w:tblLayout')
                            tblLayout.set(qn('w:type'), 'fixed')
                            tblPr.append(tblLayout)

                            # No cell spacing
                            tblCellSpacing = OxmlElement('w:tblCellSpacing')
                            tblCellSpacing.set(qn('w:w'), '0')
                            tblCellSpacing.set(qn('w:type'), 'dxa')
                            tblPr.append(tblCellSpacing)

                            tbl.append(tblPr)

                            # Table grid - 3 columns with equal width
                            tblGrid = OxmlElement('w:tblGrid')
                            for _ in range(3):
                                gridCol = OxmlElement('w:gridCol')
                                tblGrid.append(gridCol)
                            tbl.append(tblGrid)

                            # Create the single row
                            tr = OxmlElement('w:tr')

                            # Row height auto
                            trPr = OxmlElement('w:trPr')
                            trHeight = OxmlElement('w:trHeight')
                            trHeight.set(qn('w:val'), '0')
                            trHeight.set(qn('w:hRule'), 'auto')
                            trPr.append(trHeight)
                            tr.append(trPr)

                            # Cell 1: Empty (for balance) - 5%
                            tc1 = OxmlElement('w:tc')
                            tcPr1 = OxmlElement('w:tcPr')
                            tcW1 = OxmlElement('w:tcW')
                            tcW1.set(qn('w:type'), 'pct')
                            tcW1.set(qn('w:w'), '250')  # 5% (5000/100*5 = 250)
                            tcPr1.append(tcW1)
                            tc1.append(tcPr1)
                            p1 = OxmlElement('w:p')
                            tc1.append(p1)
                            tr.append(tc1)

                            # Cell 2: Formula (centered) - 80%
                            tc2 = OxmlElement('w:tc')
                            tcPr2 = OxmlElement('w:tcPr')
                            tcW2 = OxmlElement('w:tcW')
                            tcW2.set(qn('w:type'), 'pct')
                            tcW2.set(qn('w:w'), '4000')  # 80% (5000/100*80 = 4000)
                            tcPr2.append(tcW2)
                            tc2.append(tcPr2)

                            p2 = OxmlElement('w:p')
                            # Center alignment for formula
                            pPr2 = OxmlElement('w:pPr')
                            jc2 = OxmlElement('w:jc')
                            jc2.set(qn('w:val'), 'center')
                            pPr2.append(jc2)
                            p2.append(pPr2)
                            # Add the math run to this paragraph
                            p2.append(math_run)
                            tc2.append(p2)
                            tr.append(tc2)

                            # Cell 3: Equation number (right-aligned) - 15%
                            tc3 = OxmlElement('w:tc')
                            tcPr3 = OxmlElement('w:tcPr')
                            tcW3 = OxmlElement('w:tcW')
                            tcW3.set(qn('w:type'), 'pct')
                            tcW3.set(qn('w:w'), '750')  # 15% (5000/100*15 = 750)
                            tcPr3.append(tcW3)
                            tc3.append(tcPr3)

                            p3 = OxmlElement('w:p')
                            if equation_number:
                                # Right alignment for equation number
                                pPr3 = OxmlElement('w:pPr')
                                jc3 = OxmlElement('w:jc')
                                jc3.set(qn('w:val'), 'right')
                                pPr3.append(jc3)
                                p3.append(pPr3)

                                # Add the equation number as text run
                                r3 = OxmlElement('w:r')
                                rPr3 = OxmlElement('w:rPr')
                                sz3 = OxmlElement('w:sz')
                                sz3.set(qn('w:val'), '22')  # 11pt
                                rPr3.append(sz3)
                                r3.append(rPr3)

                                t3 = OxmlElement('w:t')
                                t3.text = equation_number
                                r3.append(t3)
                                p3.append(r3)

                                print(f"    ✓ Added equation number: {equation_number}")
                            tc3.append(p3)
                            tr.append(tc3)

                            # Add row to table
                            tbl.append(tr)

                            # Now insert the table into the document
                            parent = para._element.getparent()
                            para_index = parent.index(para._element)

                            if text_before:
                                print(f"    [Info] Preserving text before formula: {text_before[:50]}...")
                                # Keep the paragraph with text_before
                                para.clear()
                                para.add_run(text_before)
                                # Insert table after this paragraph
                                parent.insert(para_index + 1, tbl)
                            else:
                                # Replace the paragraph with the table
                                para.clear()
                                parent.insert(para_index, tbl)
                                parent.remove(para._element)

                            # If there's text after the placeholder, add it as a new paragraph after the table
                            if text_after:
                                print(f"    [Info] Preserving text after formula: {text_after[:50]}...")
                                new_para = OxmlElement('w:p')
                                new_run = OxmlElement('w:r')
                                new_text = OxmlElement('w:t')
                                new_text.text = text_after
                                new_run.append(new_text)
                                new_para.append(new_run)

                                # Insert after table
                                if text_before:
                                    # Table was inserted at para_index + 1, so insert text_after at para_index + 2
                                    parent.insert(para_index + 2, new_para)
                                else:
                                    # Table was inserted at para_index, so insert text_after at para_index + 1
                                    parent.insert(para_index + 1, new_para)

                            replaced_count += 1
                            print(f"    ✓ Replaced with: {latex[:60]}...")
                        else:
                            print(f"    ✗ Failed to create OMML for: {latex[:60]}...")

                    except Exception as e:
                        print(f"    ✗ Error replacing formula: {e}")
                        import traceback
                        traceback.print_exc()

                    break  # Move to next paragraph

        print(f"[Formula] Replaced {replaced_count} out of {len(formula_map)} formulas")

        if replaced_count < len(formula_map):
            print(f"[Warning] {len(formula_map) - replaced_count} placeholders not found in document")

        # Save document
        doc.save(docx_path)

        # Post-process the DOCX file to remove whitespace from formulas
        self._cleanup_formula_whitespace(docx_path)

        # Process inline math symbols in text (always enabled when formulas are detected)
        self._process_inline_math(docx_path)

    def _process_inline_math(self, docx_path):
        """
        Process inline math symbols and expressions in paragraph text and convert to OMML

        Args:
            docx_path: Path to DOCX file
        """
        from docx import Document
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import qn
        import re

        print("\n[InlineMath] Processing inline math symbols and expressions...")

        doc = Document(docx_path)

        # Define math symbol patterns
        math_symbols = {
            '∈': ('element-of', 'in'),
            '∉': ('not-element-of', 'notin'),
            '⊂': ('subset', 'subset'),
            '⊃': ('superset', 'supset'),
            '∩': ('intersection', 'cap'),
            '∪': ('union', 'cup'),
            '∀': ('for-all', 'forall'),
            '∃': ('exists', 'exists'),
            '∇': ('nabla', 'nabla'),
            '∂': ('partial', 'partial'),
            '∫': ('integral', 'int'),
            '∑': ('sum', 'sum'),
            '∏': ('product', 'prod'),
            '√': ('square-root', 'sqrt'),
            '±': ('plus-minus', 'pm'),
            '×': ('times', 'times'),
            '÷': ('divide', 'div'),
            '≤': ('less-equal', 'leq'),
            '≥': ('greater-equal', 'geq'),
            '≠': ('not-equal', 'neq'),
            '≈': ('approximately', 'approx'),
            '∞': ('infinity', 'infty'),
            '∼': ('tilde', 'sim'),
            '→': ('arrow-right', 'rightarrow'),
            '←': ('arrow-left', 'leftarrow'),
            # Greek letters
            'α': ('alpha', 'alpha'),
            'β': ('beta', 'beta'),
            'γ': ('gamma', 'gamma'),
            'δ': ('delta', 'delta'),
            'ε': ('epsilon', 'epsilon'),
            'ζ': ('zeta', 'zeta'),
            'η': ('eta', 'eta'),
            'θ': ('theta', 'theta'),
            'ι': ('iota', 'iota'),
            'κ': ('kappa', 'kappa'),
            'λ': ('lambda', 'lambda'),
            'μ': ('mu', 'mu'),
            'ν': ('nu', 'nu'),
            'ξ': ('xi', 'xi'),
            'ο': ('omicron', 'omicron'),
            'π': ('pi', 'pi'),
            'ρ': ('rho', 'rho'),
            'σ': ('sigma', 'sigma'),
            'τ': ('tau', 'tau'),
            'υ': ('upsilon', 'upsilon'),
            'φ': ('phi', 'phi'),
            'χ': ('chi', 'chi'),
            'ψ': ('psi', 'psi'),
            'ω': ('omega', 'omega'),
        }

        # Count how many paragraphs we process
        processed_count = 0
        expression_count = 0
        subscript_count = 0
        symbol_count = 0

        for para in doc.paragraphs:
            text = para.text

            # Check if paragraph contains any math symbols or expressions
            has_math = any(symbol in text for symbol in math_symbols.keys())

            # Also check for math expression patterns
            has_expressions = bool(re.search(r'[a-zA-Z]+\([^)]*[θαβγδεζηικλμνξοπρστυφχψω∈×∼][^)]*\)', text))

            # Check for subscript patterns
            has_subscripts = bool(re.search(r'\b[a-zA-Z][0-9t]\b', text))

            if has_math or has_expressions or has_subscripts:
                # Process this paragraph
                expr_count, sym_count = self._convert_inline_math_in_paragraph(para, math_symbols)
                # expr_count now includes both expressions and subscripts
                # We'll count subscripts separately in the detailed output
                expression_count += expr_count
                symbol_count += sym_count
                processed_count += 1

        print(f"[InlineMath] Processed {processed_count} paragraphs")
        print(f"[InlineMath] - {expression_count} math expressions (including subscripts)")
        print(f"[InlineMath] - {symbol_count} individual symbols")

        # Save document
        doc.save(docx_path)

    def _convert_inline_math_in_paragraph(self, para, math_symbols):
        """
        Convert inline math symbols and expressions in a paragraph to OMML format

        Args:
            para: docx.paragraph.Paragraph object
            math_symbols: Dictionary of symbol -> (name, latex) mappings

        Returns:
            tuple: (expression_count, symbol_count)
        """
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import qn
        import re

        # Get paragraph text
        text = para.text

        math_items = []  # List of (start_pos, end_pos, text, type)

        # Step 1: Find all complete math expressions
        # Define multiple expression patterns

        # Pattern 1a: Distribution expressions like "x0 ∼N(0, I)"
        # Match: variable + ∼ + function with parentheses
        dist_func_pattern = r'\b[a-zA-Z][a-zA-Z0-9]*\s*∼\s*[A-Z][a-zA-Z]*\([^)]{0,50}\)'
        for match in re.finditer(dist_func_pattern, text):
            start, end = match.span()
            expr_text = match.group()
            math_items.append((start, end, expr_text, 'expression'))

        # Pattern 1b: Set membership with ranges like "t ∈[0, 1]"
        # Match: variable + ∈ + [range]
        set_range_pattern = r'\b[a-zA-Z][a-zA-Z0-9]*\s*∈\s*\[[^\]]{0,50}\]'
        for match in re.finditer(set_range_pattern, text):
            start, end = match.span()
            expr_text = match.group()
            # Check if not already covered
            if not any(s <= start < e for s, e, _, _ in math_items):
                math_items.append((start, end, expr_text, 'expression'))

        # Pattern 1c: Set membership with space like "x ∈R1+T/4×H/8×W/8"
        # Match: variable + ∈ + complex expression with math symbols
        set_expr_pattern = r'\b[a-zA-Z][a-zA-Z0-9]*\s*∈\s*[A-Z][A-Za-z0-9\+\-\*/×÷\{\}]{2,60}'
        for match in re.finditer(set_expr_pattern, text):
            start, end = match.span()
            expr_text = match.group()
            # Check if not already covered
            if not any(s <= start < e for s, e, _, _ in math_items):
                # Trim at common delimiters
                expr_text = re.split(r'[,;\.]\s', expr_text)[0]
                end = start + len(expr_text)
                math_items.append((start, end, expr_text, 'expression'))

        # Pattern 2: Function calls with math symbols like "u(xt, ctxt, t; θ)"
        func_pattern = r'\b[a-zA-Z_]+\([^)]{0,100}[θαβγδεζηικλμνξοπρστυφχψω∈×∼∀∃∇∂∫∑∏±÷≤≥≠≈∞→←][^)]{0,100}\)'
        for match in re.finditer(func_pattern, text):
            start, end = match.span()
            expr_text = match.group()
            # Check if not already covered
            if not any(s <= start < e for s, e, _, _ in math_items):
                math_items.append((start, end, expr_text, 'expression'))

        # Pattern 3: Equations with math symbols like "xt = tx1 + (1-t)x0"
        # Match: variable = expression with math symbols (but stop at sentence boundaries)
        eq_pattern = r'\b[a-zA-Z][a-zA-Z0-9]*\s*=\s*[^\.,;]{1,80}[θαβγδεζηικλμνξοπρστυφχψω∈×∼∀∃∇∂∫∑∏±÷≤≥≠≈∞→←−][^\.,;]{0,20}'
        for match in re.finditer(eq_pattern, text):
            start, end = match.span()
            expr_text = match.group()
            # Check if not already covered and not too long
            if not any(s <= start < e for s, e, _, _ in math_items) and len(expr_text) < 100:
                # Trim at sentence boundaries or common words
                expr_text = re.split(r'[,;\.]\s+(?=[A-Z])', expr_text)[0]  # Stop at ". Word"

                # Stop before common connecting words
                common_words = r'\b(where|and|or|is|are|was|were|represents|denotes|indicates|means|such|that|which|with|for|from|to|in|on|at|by|as)\b'
                match_word = re.search(common_words, expr_text)
                if match_word:
                    expr_text = expr_text[:match_word.start()].rstrip()

                expr_text = expr_text.rstrip('.,;: ')
                end = start + len(expr_text)
                if len(expr_text) > 0:
                    math_items.append((start, end, expr_text, 'expression'))

        # Step 2: Find subscript patterns like x0, x1, xt, ctxt
        # Pattern: single letter + digits/letters (subscript)
        # Be careful not to match technical terms like HTML5, T5-small, etc.

        # Pattern 2a: Single letter + numbers (x0, x1, x2)
        subscript_pattern = r'\b([a-zA-Zα-ωΑ-Ω])([0-9]+)\b'
        for match in re.finditer(subscript_pattern, text):
            start, end = match.span()
            # Check if not already covered
            if not any(s <= start < e for s, e, _, _ in math_items):
                base = match.group(1)
                subscript = match.group(2)

                # Skip common technical terms
                full_text = match.group(0)
                if full_text.lower() in ['t5', 'h264', 'h265', 'utf8', 'base64', 'md5', 'sha256']:
                    continue

                math_items.append((start, end, f"{base}_{subscript}", 'subscript'))

        # Pattern 2b: Single letter + letter (xt means x_t, vt means v_t)
        subscript_letter_pattern = r'\b([a-zA-Z])([a-z])\b'
        for match in re.finditer(subscript_letter_pattern, text):
            start, end = match.span()
            # Check if not already covered
            if not any(s <= start < e for s, e, _, _ in math_items):
                base = match.group(1)
                subscript = match.group(2)
                full_text = match.group(0)

                # Only if both are lowercase or first is uppercase
                # Skip common words like "is", "it", "at", "in", "on", "to", etc.
                if full_text in ['is', 'it', 'at', 'in', 'on', 'to', 'of', 'or', 'an', 'as', 'be', 'by', 'he', 'we', 'if', 'so', 'up', 'no']:
                    continue

                # Only if looks like a variable (short and mathematical context)
                if len(full_text) == 2:
                    math_items.append((start, end, f"{base}_{subscript}", 'subscript'))

        # Pattern 2c: Multi-letter variable + letter subscript (ctxt means c_{txt})
        multi_subscript_pattern = r'\b([a-z]{1,3})([a-z]{1,3})\b'
        for match in re.finditer(multi_subscript_pattern, text):
            start, end = match.span()
            full_text = match.group(0)

            # Only process specific known patterns from the context
            if full_text in ['ctxt']:  # Add more as needed
                # Check if not already covered
                if not any(s <= start < e for s, e, _, _ in math_items):
                    base = match.group(1)
                    subscript = match.group(2)
                    math_items.append((start, end, f"{base}_{subscript}", 'subscript'))

        # Step 3: Find individual math symbols (not already in expressions)
        for symbol in math_symbols.keys():
            pos = 0
            while True:
                pos = text.find(symbol, pos)
                if pos == -1:
                    break

                # Check if this symbol is already part of an expression
                in_expression = False
                for item_start, item_end, _, item_type in math_items:
                    if item_start <= pos < item_end:
                        in_expression = True
                        break

                if not in_expression:
                    math_items.append((pos, pos + len(symbol), symbol, 'symbol'))

                pos += 1

        # If no math items, return
        if not math_items:
            return 0, 0

        # Sort by position
        math_items.sort(key=lambda x: x[0])

        # Remove overlapping items (keep the first one)
        filtered_items = []
        last_end = -1
        for start, end, content, item_type in math_items:
            if start >= last_end:
                filtered_items.append((start, end, content, item_type))
                last_end = end

        # Count items
        expression_count = sum(1 for _, _, _, t in filtered_items if t == 'expression')
        subscript_count = sum(1 for _, _, _, t in filtered_items if t == 'subscript')
        symbol_count = sum(1 for _, _, _, t in filtered_items if t == 'symbol')

        # Build new paragraph content with math runs
        parent = para._element.getparent()
        para_index = parent.index(para._element)

        # Create new paragraph element
        new_para = OxmlElement('w:p')

        # Copy paragraph properties if any
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            new_para.append(pPr)

        # Build runs: text segments + math runs
        last_pos = 0
        for start, end, content, item_type in filtered_items:
            # Add text before this math item
            if start > last_pos:
                text_before = text[last_pos:start]
                text_run = OxmlElement('w:r')
                text_elem = OxmlElement('w:t')
                text_elem.set(qn('xml:space'), 'preserve')
                text_elem.text = text_before
                text_run.append(text_elem)
                new_para.append(text_run)

            # Add math run
            if item_type == 'expression':
                math_run = self._create_inline_math_expression(content)
            elif item_type == 'subscript':
                math_run = self._create_subscript_math(content)
            else:  # symbol
                math_run = self._create_inline_math_run(content)

            if math_run is not None:
                new_para.append(math_run)

            last_pos = end

        # Add remaining text
        if last_pos < len(text):
            text_after = text[last_pos:]
            text_run = OxmlElement('w:r')
            text_elem = OxmlElement('w:t')
            text_elem.set(qn('xml:space'), 'preserve')
            text_elem.text = text_after
            text_run.append(text_elem)
            new_para.append(text_run)

        # Replace old paragraph with new one
        parent.insert(para_index, new_para)
        parent.remove(para._element)

        return expression_count + subscript_count, symbol_count

    def _create_inline_math_expression(self, expr_text):
        """
        Create an inline math run (OMML) for a complete expression like u(xt, ctxt, t; θ)

        Args:
            expr_text: Expression string like "u(xt, ctxt, t; θ)"

        Returns:
            OxmlElement (w:r with m:oMath) or None
        """
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import qn
        import re

        try:
            # Build OMML content by processing the expression text
            # We need to identify and convert subscripts within the expression

            omml_parts = []
            i = 0
            text_len = len(expr_text)

            while i < text_len:
                # Try to match subscript patterns at current position
                remaining = expr_text[i:]

                # Pattern 1: Multi-letter + multi-letter (ctxt -> c_{txt})
                # Must check this FIRST before single letter patterns
                match = re.match(r'^(c)(txt)(?![a-z])', remaining)
                if match:
                    base = match.group(1)
                    subscript = match.group(2)
                    # Create subscript OMML
                    omml_parts.append(
                        f'<m:sSub>'
                        f'<m:e><m:r><m:t>{base}</m:t></m:r></m:e>'
                        f'<m:sub><m:r><m:t>{subscript}</m:t></m:r></m:sub>'
                        f'</m:sSub>'
                    )
                    i += len(match.group(0))
                    continue

                # Pattern 2: Single letter + digits (x0, x1)
                match = re.match(r'^([a-zA-Zα-ωΑ-Ω])([0-9]+)', remaining)
                if match:
                    base = match.group(1)
                    subscript = match.group(2)
                    # Create subscript OMML
                    omml_parts.append(
                        f'<m:sSub>'
                        f'<m:e><m:r><m:t>{base}</m:t></m:r></m:e>'
                        f'<m:sub><m:r><m:t>{subscript}</m:t></m:r></m:sub>'
                        f'</m:sSub>'
                    )
                    i += len(match.group(0))
                    continue

                # Pattern 3: Single letter + single letter (xt, vt)
                # But be careful with common words
                match = re.match(r'^([a-zA-Z])([a-z])(?![a-z])', remaining)
                if match:
                    base = match.group(1)
                    subscript = match.group(2)
                    full = match.group(0)

                    # Skip common words
                    if full not in ['is', 'it', 'at', 'in', 'on', 'to', 'of', 'or', 'an', 'as', 'be', 'by']:
                        # Create subscript OMML
                        omml_parts.append(
                            f'<m:sSub>'
                            f'<m:e><m:r><m:t>{base}</m:t></m:r></m:e>'
                            f'<m:sub><m:r><m:t>{subscript}</m:t></m:r></m:sub>'
                            f'</m:sSub>'
                        )
                        i += len(match.group(0))
                        continue

                # No subscript pattern matched, add single character as text
                char = expr_text[i]
                omml_parts.append(f'<m:r><m:t>{char}</m:t></m:r>')
                i += 1

            # Combine all parts
            omml_content = ''.join(omml_parts)

            # Remove whitespace between tags
            omml_content = re.sub(r'>\s+<', '><', omml_content)

            # Create OMML run
            omml_xml = (
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                f'<m:oMath>{omml_content}</m:oMath>'
                '</w:r>'
            )

            omml_elem = parse_xml(omml_xml)
            return omml_elem

        except Exception as e:
            print(f"[Warning] Failed to create inline math expression for '{expr_text}': {e}")
            import traceback
            traceback.print_exc()
            return None

    def _create_subscript_math(self, subscript_text):
        """
        Create an inline math run (OMML) for subscript like x_0, v_t, c_{txt}

        Args:
            subscript_text: Subscript string in format "base_subscript" (e.g., "x_0", "v_t")

        Returns:
            OxmlElement (w:r with m:oMath containing m:sSub) or None
        """
        from docx.oxml import parse_xml
        import re

        try:
            # Parse the subscript notation
            parts = subscript_text.split('_')
            if len(parts) != 2:
                return self._create_inline_math_run(subscript_text)

            base, subscript = parts

            # Build OMML subscript structure
            # <m:sSub> = subscript element
            #   <m:e> = base element
            #   <m:sub> = subscript element

            omml_content = (
                f'<m:sSub>'
                f'<m:e><m:r><m:t>{base}</m:t></m:r></m:e>'
                f'<m:sub><m:r><m:t>{subscript}</m:t></m:r></m:sub>'
                f'</m:sSub>'
            )

            # Remove whitespace between tags
            omml_content = re.sub(r'>\s+<', '><', omml_content)

            # Create OMML run
            omml_xml = (
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                f'<m:oMath>{omml_content}</m:oMath>'
                '</w:r>'
            )

            omml_elem = parse_xml(omml_xml)
            return omml_elem

        except Exception as e:
            print(f"[Warning] Failed to create subscript math for '{subscript_text}': {e}")
            return None

    def _create_inline_math_run(self, symbol):
        """
        Create an inline math run (OMML) for a single symbol

        Args:
            symbol: Unicode math symbol string

        Returns:
            OxmlElement (w:r with m:oMath) or None
        """
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import qn

        try:
            # Create a simple OMML structure for the symbol
            # We'll use m:r (math run) with m:t (math text)
            omml_xml = (
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                '<m:oMath><m:r><m:t>' + symbol + '</m:t></m:r></m:oMath>'
                '</w:r>'
            )

            omml_elem = parse_xml(omml_xml)
            return omml_elem

        except Exception as e:
            print(f"[Warning] Failed to create inline math for symbol '{symbol}': {e}")
            return None

    def _cleanup_formula_whitespace(self, docx_path):
        """
        Post-process DOCX file to remove whitespace from math formulas

        Args:
            docx_path: Path to DOCX file
        """
        import zipfile
        import tempfile
        import os
        import re
        from lxml import etree

        try:
            # Create a temporary file
            temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_fd)

            # Open the DOCX as a ZIP file
            with zipfile.ZipFile(docx_path, 'r') as zip_in:
                with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                    for item in zip_in.namelist():
                        data = zip_in.read(item)

                        # Process word/document.xml to clean formula whitespace
                        if item == 'word/document.xml':
                            # Parse XML
                            root = etree.fromstring(data)

                            # Find all m:oMath elements and remove whitespace
                            nsmap = {
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                                'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                            }

                            for math_elem in root.xpath('.//m:oMath', namespaces=nsmap):
                                # Remove text content (whitespace) directly under m:oMath
                                if math_elem.text and not math_elem.text.strip():
                                    math_elem.text = None

                                # Remove whitespace tails from all descendants
                                for elem in math_elem.iter():
                                    if elem.tail and not elem.tail.strip():
                                        elem.tail = None

                            # Serialize back to bytes without pretty printing
                            data = etree.tostring(root, encoding='utf-8', xml_declaration=True)

                        # Write to new ZIP
                        zip_out.writestr(item, data)

            # Replace original file with cleaned version
            import shutil
            shutil.move(temp_path, docx_path)

            print(f"[Formula] Cleaned whitespace from formulas")

        except Exception as e:
            print(f"[Warning] Failed to cleanup formula whitespace: {e}")
            import traceback
            traceback.print_exc()
            # Clean up temp file if it exists
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass

    def _remove_blank_text_nodes(self, element):
        """
        Recursively remove blank text nodes from an XML element tree

        Args:
            element: lxml Element
        """
        # Remove blank text at element start
        if element.text and not element.text.strip():
            element.text = None

        # Process all children
        for child in element:
            # Remove blank tail of previous sibling
            if child.tail and not child.tail.strip():
                child.tail = None
            # Recursively process child
            self._remove_blank_text_nodes(child)

    def _mathml_to_omml(self, mathml_elem):
        """
        Convert MathML element to OMML (Office Math ML) element recursively

        Args:
            mathml_elem: lxml Element with MathML

        Returns:
            OMML element as XML string
        """
        from lxml import etree

        # Remove namespace for easier handling
        tag = mathml_elem.tag.split('}')[-1] if '}' in mathml_elem.tag else mathml_elem.tag

        # Handle different MathML elements
        if tag == 'math':
            # Root element - process children
            omml_parts = []
            for child in mathml_elem:
                omml_parts.append(self._mathml_to_omml(child))
            return ''.join(omml_parts)

        elif tag == 'mrow':
            # Group - just process children
            omml_parts = []
            for child in mathml_elem:
                omml_parts.append(self._mathml_to_omml(child))
            return ''.join(omml_parts)

        elif tag == 'msup':
            # Superscript: base^superscript
            children = list(mathml_elem)
            if len(children) >= 2:
                base = self._mathml_to_omml(children[0])
                sup = self._mathml_to_omml(children[1])
                return f'<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'
            return ''

        elif tag == 'msub':
            # Subscript: base_subscript
            children = list(mathml_elem)
            if len(children) >= 2:
                base = self._mathml_to_omml(children[0])
                sub = self._mathml_to_omml(children[1])
                return f'<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'
            return ''

        elif tag == 'msubsup':
            # Subscript and superscript: base_sub^sup
            children = list(mathml_elem)
            if len(children) >= 3:
                base = self._mathml_to_omml(children[0])
                sub = self._mathml_to_omml(children[1])
                sup = self._mathml_to_omml(children[2])
                return f'<m:sSubSup><m:e>{base}</m:e><m:sub>{sub}</m:sub><m:sup>{sup}</m:sup></m:sSubSup>'
            return ''

        elif tag == 'mfrac':
            # Fraction: numerator/denominator
            children = list(mathml_elem)
            if len(children) >= 2:
                num = self._mathml_to_omml(children[0])
                den = self._mathml_to_omml(children[1])
                return f'<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'
            return ''

        elif tag in ['mi', 'mn', 'mo', 'mtext']:
            # Text elements: identifier, number, operator, text
            text = mathml_elem.text or ''
            # Escape XML special characters
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return f'<m:r><m:t>{text}</m:t></m:r>'

        else:
            # Unknown element - just process children and extract text
            text = ''.join(mathml_elem.itertext())
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return f'<m:r><m:t>{text}</m:t></m:r>'

    def _create_omml_from_mathml(self, mathml):
        """
        Create OMML (Office Math ML) run element from MathML

        Args:
            mathml: MathML string

        Returns:
            OxmlElement (w:r with m:oMath) or None
        """
        from docx.oxml import parse_xml
        from lxml import etree
        import re

        try:
            # Parse the MathML
            mathml_tree = etree.fromstring(mathml.encode('utf-8'))

            # Convert MathML to OMML
            omml_content = self._mathml_to_omml(mathml_tree).strip()

            # Remove all inter-tag whitespace (between > and <)
            omml_content = re.sub(r'>\s+<', '><', omml_content)

            # Create OMML run with math zone (no whitespace!)
            omml_xml = (
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                f'<m:oMath>{omml_content}</m:oMath>'
                '</w:r>'
            )

            # Parse with lxml first to clean whitespace
            parser = etree.XMLParser(remove_blank_text=True)
            lxml_elem = etree.fromstring(omml_xml.encode('utf-8'), parser=parser)

            # Convert back to string without pretty printing
            clean_xml = etree.tostring(lxml_elem, encoding='unicode')

            # Now parse with docx parser
            omml_elem = parse_xml(clean_xml)
            return omml_elem

        except Exception as e:
            print(f"[Warning] Failed to create OMML: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _inject_yolo_to_pdf2docx(self, pdf_path):
        """
        Inject YOLO detected regions into pdf2docx's image extraction
        Also filter out text/shapes that overlap with YOLO detected regions

        This modifies pdf2docx's internal image extraction to use YOLO results
        """
        # Monkey-patch pdf2docx's ImagesExtractor
        from pdf2docx.image.ImagesExtractor import ImagesExtractor
        from pdf2docx.shape.Paths import Paths
        from pdf2docx.shape.Shapes import Shapes
        from pdf2docx.page.RawPageFitz import RawPageFitz

        # Save original methods
        original_extract_images = ImagesExtractor.extract_images
        original_paths_restore = Paths.restore
        original_shapes_restore = Shapes.restore
        original_to_shapes_and_images = Paths.to_shapes_and_images
        original_preprocess_text = RawPageFitz._preprocess_text

        yolo_detector = self.yolo_detector

        def _rect_overlap_ratio(rect1, rect2):
            """Calculate overlap ratio between two rectangles"""
            # Convert to fitz.Rect if needed
            r1 = fitz.Rect(rect1) if not isinstance(rect1, fitz.Rect) else rect1
            r2 = fitz.Rect(rect2) if not isinstance(rect2, fitz.Rect) else rect2

            # Get intersection
            intersection = r1 & r2
            if intersection.is_empty:
                return 0.0

            # Calculate areas
            intersection_area = intersection.width * intersection.height
            r1_area = r1.width * r1.height

            if r1_area == 0:
                return 0.0

            return intersection_area / r1_area

        def yolo_enhanced_extract_images(self, clip_image_res_ratio=3.0, **kwargs):
            """Enhanced image extraction using YOLO detections (figures only)"""

            page = self._page
            page_num = page.number

            # Get YOLO detected regions for this page
            regions = yolo_detector.detected_regions.get(page_num, {})
            yolo_figures = regions.get('figures', [])

            if yolo_figures:
                # Use YOLO detections for figures only
                images = []

                # Add figures
                for fig in yolo_figures:
                    bbox = fig['bbox']
                    rect = fitz.Rect(bbox)

                    # Extract image from PDF
                    mat = fitz.Matrix(clip_image_res_ratio, clip_image_res_ratio)
                    pix = page.get_pixmap(matrix=mat, clip=rect)

                    # Convert to pdf2docx format
                    if pix.colorspace.n > 3:
                        pix = fitz.Pixmap(fitz.csRGB, pix)

                    raw_dict = {
                        "type": 1,  # BlockType.IMAGE
                        "bbox": tuple(bbox),
                        "width": pix.width,
                        "height": pix.height,
                        "image": pix.tobytes(),
                    }

                    images.append(raw_dict)

                print(f"[YOLO] Page {page_num + 1}: Using {len(yolo_figures)} YOLO-detected figures")
                return images
            else:
                # Fall back to original pdf2docx image extraction
                result = original_extract_images(self, clip_image_res_ratio, **kwargs)
                if result:
                    print(f"[Debug] Page {page_num + 1}: extract_images() returned {len(result)} images (no YOLO)")
                return result

        def yolo_filtered_preprocess_text(self, **settings):
            """Filter out text blocks that overlap with YOLO regions, but insert formula placeholders"""
            # Get page number
            page_num = self.page_engine.number
            regions = yolo_detector.detected_regions.get(page_num, {})
            yolo_figures = regions.get('figures', [])

            # Call original method to get all text blocks
            text_blocks = original_preprocess_text(self, **settings)

            # If no YOLO figures or formulas, return original result
            yolo_formulas = regions.get('formulas', [])
            if not yolo_figures and not yolo_formulas:
                return text_blocks

            # First, detect equation numbers for formulas BEFORE filtering
            # Extract directly from PDF page instead of relying on text_blocks
            page = self.page_engine
            for formula_idx, formula in enumerate(yolo_formulas):
                bbox = formula['bbox']
                formula_rect = fitz.Rect(bbox)

                # Debug: print formula bbox
                print(f"[Debug] Page {page_num + 1}, Formula {formula_idx} bbox: {bbox}")

                # Try to find equation number by extracting text from area to the right of formula
                # Create a search rectangle to the right of the formula
                search_rect = fitz.Rect(
                    formula_rect.x1 - 5,  # Start slightly before formula end
                    formula_rect.y0 - 5,  # Expand vertically a bit
                    formula_rect.x1 + 100,  # Search 100 points to the right
                    formula_rect.y1 + 5
                )

                # Extract text from this region
                text_dict = page.get_textpage().extractDICT()
                equation_number = None

                for block in text_dict.get('blocks', []):
                    if block.get('type') != 0:  # Only text blocks
                        continue

                    block_bbox = fitz.Rect(block['bbox'])

                    # Check if this block intersects with our search rectangle
                    if not search_rect.intersects(block_bbox):
                        continue

                    # Extract text from this block
                    block_text = ''
                    for line in block.get('lines', []):
                        for span in line.get('spans', []):
                            block_text += span.get('text', '')

                    if not block_text.strip():
                        continue

                    print(f"[Debug]   Found text in search area: '{block_text}'")

                    # Check if it matches equation number pattern
                    import re
                    match = re.search(r'\(\s*\d+(\.\d+)?\s*\)', block_text)
                    if match:
                        equation_number = match.group(0)
                        print(f"[Formula] Detected equation number: '{equation_number}' for formula {formula_idx}")
                        break

                # Store in formula metadata
                if equation_number:
                    formula['equation_number'] = equation_number

            # Convert YOLO figure regions to Rect list (NOT formulas - we want to keep formula text areas)
            yolo_rects = [fitz.Rect(fig['bbox']) for fig in yolo_figures]

            # Filter out text blocks that overlap with YOLO figure regions (but not formula regions)
            # IMPROVED: Use higher threshold to preserve body text near figures
            filtered_blocks = []
            filtered_count = 0
            for block in text_blocks:
                # Check if this is a text block (type 0) vs image block (type 1)
                if block.get('type') != 0:
                    filtered_blocks.append(block)
                    continue

                block_bbox = fitz.Rect(block['bbox'])
                overlaps_with_figure = False
                for yolo_rect in yolo_rects:
                    overlap_ratio = _rect_overlap_ratio(block_bbox, yolo_rect)
                    # IMPROVED: Raised threshold from 0.5 to 0.85
                    # Only filter text that's almost completely inside a figure
                    # This preserves body text that's merely adjacent to figures
                    if overlap_ratio > 0.85:  # More than 85% overlap with figure
                        overlaps_with_figure = True
                        filtered_count += 1
                        break

                if not overlaps_with_figure:
                    filtered_blocks.append(block)

            # Insert formula placeholder text blocks
            for formula_idx, formula in enumerate(yolo_formulas):
                bbox = formula['bbox']
                formula_rect = fitz.Rect(bbox)
                placeholder_text = f"FORMULA_PLACEHOLDER_{page_num}_{formula_idx}"

                # Get equation number if detected earlier
                equation_number = formula.get('equation_number', None)

                # Create a text block for the placeholder
                placeholder_block = {
                    'type': 0,  # Text block
                    'bbox': tuple(bbox),
                    'lines': [{
                        'bbox': tuple(bbox),
                        'wmode': 0,
                        'dir': [1, 0],
                        'spans': [{
                            'bbox': tuple(bbox),
                            'text': placeholder_text,
                            'size': 12,
                            'flags': 0,
                            'font': 'Arial',
                            'color': 0
                        }]
                    }]
                }

                # Add equation number as a separate span if found
                if equation_number:
                    # Extend bbox to include number position (to the right)
                    extended_bbox = list(bbox)
                    extended_bbox[2] += 50  # Add some width for the number

                    placeholder_block['lines'][0]['spans'].append({
                        'bbox': tuple(extended_bbox),
                        'text': f" {equation_number}",  # Add space before number
                        'size': 12,
                        'flags': 0,
                        'font': 'Arial',
                        'color': 0
                    })

                filtered_blocks.append(placeholder_block)

            if filtered_count > 0:
                print(f"[Filter] Page {page_num + 1}: Filtered {filtered_count} text blocks overlapping with figures")
            if yolo_formulas:
                print(f"[Filter] Page {page_num + 1}: Inserted {len(yolo_formulas)} formula placeholders")

            return filtered_blocks

        def yolo_filtered_paths_restore(self, raws:list):
            """Completely skip paths for pages with YOLO detections"""
            # Get page number
            page_num = getattr(self.parent, 'number', 0)
            regions = yolo_detector.detected_regions.get(page_num, {})
            yolo_figures = regions.get('figures', [])
            yolo_formulas = regions.get('formulas', [])

            # If page has YOLO detections (figures or formulas), skip all path processing
            if yolo_figures or yolo_formulas:
                print(f"[Filter] Page {page_num + 1}: Disabled all path processing (using YOLO only)")
                return self  # Return empty paths
            else:
                # No YOLO detections, use original method
                return original_paths_restore(self, raws)

        def yolo_filtered_shapes_restore(self, raws:list):
            """Filter out shapes that overlap with YOLO detected regions"""
            # Get page number
            page_num = getattr(self.parent, 'number', 0)
            regions = yolo_detector.detected_regions.get(page_num, {})
            yolo_figures = regions.get('figures', [])
            yolo_formulas = regions.get('formulas', [])

            # If no YOLO figures or formulas, use original method
            if not yolo_figures and not yolo_formulas:
                return original_shapes_restore(self, raws)

            # Convert YOLO regions to Rect list (both figures and formulas)
            yolo_rects = [fitz.Rect(fig['bbox']) for fig in yolo_figures]
            yolo_rects.extend([fitz.Rect(formula['bbox']) for formula in yolo_formulas])

            # Filter shapes
            from pdf2docx.shape.Shape import Stroke, Fill, Hyperlink
            self.reset()
            rect = (0, 0, self.parent.width, self.parent.height)
            filtered_count = 0

            for raw in raws:
                # Distinguish specified type by key like `start`, `end` and `uri` (same as original)
                if 'start' in raw:
                    shape = Stroke(raw)
                elif 'uri' in raw:
                    shape = Hyperlink(raw)
                else:
                    shape = Fill(raw)

                # Ignore shape out of page
                if not shape.bbox.intersects(rect):
                    continue

                # Check if shape overlaps with any YOLO region
                overlaps = False
                for yolo_rect in yolo_rects:
                    overlap_ratio = _rect_overlap_ratio(shape.bbox, yolo_rect)
                    # IMPROVED: Keep at 90% for shapes (more conservative than text)
                    if overlap_ratio > 0.9:  # More than 90% overlap
                        overlaps = True
                        filtered_count += 1
                        break

                if not overlaps:
                    self.append(shape)

            if filtered_count > 0:
                print(f"[Filter] Page {page_num + 1}: Filtered {filtered_count} overlapping shapes")

            return self

        def yolo_filtered_to_shapes_and_images(self,
                                               min_svg_gap_dx:float=15,
                                               min_svg_gap_dy:float=15,
                                               min_w:float=2,
                                               min_h:float=2,
                                               clip_image_res_ratio:float=3.0):
            """Completely disable SVG image generation for pages with YOLO detections"""
            # Get page number
            page_num = getattr(self.parent, 'number', 0)
            regions = yolo_detector.detected_regions.get(page_num, {})
            yolo_figures = regions.get('figures', [])
            yolo_formulas = regions.get('formulas', [])

            # If page has YOLO detections (figures or formulas), disable all SVG/Path image generation
            if yolo_figures or yolo_formulas:
                # Only generate shapes (for tables/text styles), NO images
                iso_shapes = []
                if self.is_iso_oriented:
                    iso_shapes.extend(self.to_shapes())

                print(f"[Filter] Page {page_num + 1}: Disabled SVG/Path image generation (using YOLO only)")
                return iso_shapes, []  # Return empty images list
            else:
                # No YOLO detections, use original method
                return original_to_shapes_and_images(
                    self, min_svg_gap_dx, min_svg_gap_dy, min_w, min_h, clip_image_res_ratio)

        # Apply monkey-patches
        ImagesExtractor.extract_images = yolo_enhanced_extract_images
        RawPageFitz._preprocess_text = yolo_filtered_preprocess_text
        Paths.restore = yolo_filtered_paths_restore
        Paths.to_shapes_and_images = yolo_filtered_to_shapes_and_images
        # Shapes.restore = yolo_filtered_shapes_restore  # Disabled: may interfere with table parsing


def main():
    """Test the hybrid converter"""
    import argparse

    parser = argparse.ArgumentParser(description='Hybrid PDF to DOCX Converter')
    parser.add_argument('--pdf', required=True, help='Input PDF file')
    parser.add_argument('--output', help='Output DOCX file (default: input_hybrid.docx)')
    parser.add_argument('--yolo-model', default='weights/doclayout_yolo_docstructbench_imgsz1024.pt')
    parser.add_argument('--yolo-conf', type=float, default=0.25, help='YOLO confidence threshold')
    parser.add_argument('--yolo-imgsz', type=int, default=1024, help='YOLO image size')
    parser.add_argument('--no-yolo', action='store_true', help='Disable YOLO, use pdf2docx only')

    args = parser.parse_args()

    # Set default output path
    if args.output is None:
        args.output = args.pdf.replace('.pdf', '_hybrid.docx')

    # Create converter
    converter = HybridConverter(yolo_model_path=args.yolo_model)

    # Convert
    converter.convert(
        pdf_path=args.pdf,
        docx_path=args.output,
        use_yolo_images=not args.no_yolo,
        yolo_conf=args.yolo_conf,
        yolo_imgsz=args.yolo_imgsz
    )


if __name__ == "__main__":
    main()
