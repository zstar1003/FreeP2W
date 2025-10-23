"""
Deep comparison between our output and smallpdf
"""

from docx import Document
import re

def deep_analyze_paragraph(para, idx, doc_name):
    """Analyze a single paragraph in detail"""
    text = para.text.strip()
    if not text or len(text) < 20:
        return None

    # Get all formatting details
    info = {
        'doc': doc_name,
        'idx': idx,
        'text': text[:80],
        'alignment': str(para.alignment).split('.')[-1] if para.alignment else 'None',
        'font_names': [],
        'font_sizes': [],
        'space_before': para.paragraph_format.space_before,
        'space_after': para.paragraph_format.space_after,
        'line_spacing': para.paragraph_format.line_spacing,
        'first_line_indent': para.paragraph_format.first_line_indent,
        'left_indent': para.paragraph_format.left_indent,
        'right_indent': para.paragraph_format.right_indent,
    }

    # Get font info from runs
    for run in para.runs:
        if run.font.name:
            info['font_names'].append(run.font.name)
        if run.font.size:
            info['font_sizes'].append(run.font.size.pt)

    return info

def compare_documents(smallpdf_path, our_path):
    """Compare two documents paragraph by paragraph"""

    print("="*80)
    print("DETAILED DOCUMENT COMPARISON")
    print("="*80)

    doc_smallpdf = Document(smallpdf_path)
    doc_ours = Document(our_path)

    # Get comparable paragraphs (body text only, skip empty and short ones)
    smallpdf_paras = []
    our_paras = []

    for i, para in enumerate(doc_smallpdf.paragraphs[:50]):
        info = deep_analyze_paragraph(para, i, 'smallpdf')
        if info:
            smallpdf_paras.append(info)

    for i, para in enumerate(doc_ours.paragraphs[:50]):
        info = deep_analyze_paragraph(para, i, 'ours')
        if info:
            our_paras.append(info)

    # Compare side by side
    print(f"\nFound {len(smallpdf_paras)} analyzable paragraphs in smallpdf")
    print(f"Found {len(our_paras)} analyzable paragraphs in our output")

    # Find matching paragraphs by text similarity
    print("\n" + "="*80)
    print("SIDE-BY-SIDE COMPARISON (First 10 matching paragraphs)")
    print("="*80)

    matched = 0
    with open('detailed_comparison.txt', 'w', encoding='utf-8') as f:
        f.write("="*80 + "\n")
        f.write("DETAILED PARAGRAPH COMPARISON\n")
        f.write("="*80 + "\n\n")

        for sp in smallpdf_paras[:20]:
            # Try to find matching paragraph in our output
            for op in our_paras:
                # Simple match by first 30 chars
                if sp['text'][:30].lower() == op['text'][:30].lower():
                    matched += 1

                    f.write(f"\n{'='*80}\n")
                    f.write(f"PARAGRAPH {matched}: {sp['text'][:60]}...\n")
                    f.write(f"{'='*80}\n\n")

                    f.write("SMALLPDF:\n")
                    f.write(f"  Alignment: {sp['alignment']}\n")
                    f.write(f"  Font names: {set(sp['font_names'])}\n")
                    f.write(f"  Font sizes: {set(sp['font_sizes'])}\n")
                    f.write(f"  Space before: {sp['space_before']}\n")
                    f.write(f"  Space after: {sp['space_after']}\n")
                    f.write(f"  Line spacing: {sp['line_spacing']}\n")
                    f.write(f"  First line indent: {sp['first_line_indent']}\n")
                    f.write(f"  Left indent: {sp['left_indent']}\n")
                    f.write(f"  Right indent: {sp['right_indent']}\n")

                    f.write("\nOUR OUTPUT:\n")
                    f.write(f"  Alignment: {op['alignment']}\n")
                    f.write(f"  Font names: {set(op['font_names'])}\n")
                    f.write(f"  Font sizes: {set(op['font_sizes'])}\n")
                    f.write(f"  Space before: {op['space_before']}\n")
                    f.write(f"  Space after: {op['space_after']}\n")
                    f.write(f"  Line spacing: {op['line_spacing']}\n")
                    f.write(f"  First line indent: {op['first_line_indent']}\n")
                    f.write(f"  Left indent: {op['left_indent']}\n")
                    f.write(f"  Right indent: {op['right_indent']}\n")

                    f.write("\nDIFFERENCES:\n")
                    if sp['alignment'] != op['alignment']:
                        f.write(f"  ⚠ Alignment: {sp['alignment']} vs {op['alignment']}\n")
                    if set(sp['font_names']) != set(op['font_names']):
                        f.write(f"  ⚠ Font names differ\n")
                    if set(sp['font_sizes']) != set(op['font_sizes']):
                        f.write(f"  ⚠ Font sizes differ\n")
                    if sp['space_before'] != op['space_before']:
                        f.write(f"  ⚠ Space before differs\n")
                    if sp['space_after'] != op['space_after']:
                        f.write(f"  ⚠ Space after differs\n")
                    if sp['line_spacing'] != op['line_spacing']:
                        f.write(f"  ⚠ Line spacing differs\n")

                    break

    print(f"\nMatched {matched} paragraphs")
    print("\nDetailed comparison saved to: detailed_comparison.txt")

    # Summary of differences
    print("\n" + "="*80)
    print("KEY DIFFERENCES SUMMARY")
    print("="*80)

    # Check overall font usage
    smallpdf_fonts = set()
    our_fonts = set()

    for sp in smallpdf_paras:
        smallpdf_fonts.update(sp['font_names'])

    for op in our_paras:
        our_fonts.update(op['font_names'])

    print(f"\nFonts in smallpdf: {smallpdf_fonts}")
    print(f"Fonts in our output: {our_fonts}")

    if smallpdf_fonts != our_fonts:
        print("\n⚠ Font mismatch detected!")

# Run comparison
compare_documents('test_files/2503.20314v2_smallpdf.docx', 'result_smallpdf_style.docx')
