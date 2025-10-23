"""
Complete PDF to DOCX conversion with smallpdf-style formatting
"""

import sys
import os
from hybrid_converter import HybridConverter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
import re

def format_to_smallpdf_style(docx_path):
    """
    Apply smallpdf-style formatting to converted DOCX
    """
    doc = Document(docx_path)

    print("\n[Format] Applying smallpdf-style formatting...")

    # Determine body text size
    font_sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                font_sizes.append(run.font.size.pt)

    body_size = Counter(font_sizes).most_common(1)[0][0] if font_sizes else 10.0
    print(f"[Format] Body text size: {body_size}pt")

    # Detect TOC
    toc_start = None
    toc_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if toc_start is None and ('contents' in text or 'table of contents' in text):
            toc_start = i

        if toc_start is not None and toc_end is None and i > toc_start:
            para_text = para.text.strip()
            if (re.match(r'^\d+\.?\d*\s+[A-Z]', para_text) or
                any(keyword in text for keyword in ['abstract', 'introduction']) or
                len(para_text) > 100):
                toc_end = i
                break

    if toc_start and not toc_end:
        toc_end = min(toc_start + 15, len(doc.paragraphs))

    # Format paragraphs
    stats = {'titles': 0, 'sections': 0, 'figures': 0, 'body': 0, 'toc': 0, 'math': 0}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Skip math paragraphs
        has_math = 'oMath' in para._element.xml if hasattr(para, '_element') else False
        if has_math:
            stats['math'] += 1
            continue

        # Detect paragraph type
        max_font_size = 0
        for run in para.runs:
            if run.font.size:
                max_font_size = max(max_font_size, run.font.size.pt)

        is_figure = text.lower().startswith(('figure', 'fig.', 'fig ', 'table'))
        is_section = bool(re.match(r'^[\d\.]+\s+[A-Z]', text))
        is_title = max_font_size > 12 or (len(text) < 100 and text.isupper() and not is_figure)
        in_toc = toc_start is not None and toc_start <= i < toc_end

        # Apply formatting
        if in_toc:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cleaned = re.sub(r'\s+', ' ', text).strip()
            if cleaned != text and para.runs:
                para.clear()
                run = para.add_run(cleaned)
                run.font.name = 'Bookman Old Style'
                run.font.size = Pt(body_size)
            stats['toc'] += 1

        elif is_figure:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Bookman Old Style'
                if run.font.size and run.font.size.pt != body_size:
                    run.font.size = Pt(body_size)
            stats['figures'] += 1

        elif is_section:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = 'Bookman Old Style'
                if run.font.size and run.font.size.pt <= 12:
                    run.font.size = Pt(body_size)
            stats['sections'] += 1

        elif is_title:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats['titles'] += 1

        else:
            # Body text: JUSTIFY alignment (smallpdf style)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = 'Bookman Old Style'
                if run.font.size and run.font.size.pt <= 12:
                    run.font.size = Pt(body_size)
            stats['body'] += 1

    print(f"[Format] Formatted {stats['sections']} sections, {stats['figures']} figures, {stats['body']} body paragraphs")

    # Save
    doc.save(docx_path)
    print(f"[Format] Formatting completed")


def convert_pdf_to_docx(pdf_path, output_path=None, apply_formatting=True):
    """
    Complete conversion pipeline:
    1. Convert PDF to DOCX using hybrid converter
    2. Apply smallpdf-style formatting
    """

    if output_path is None:
        output_path = pdf_path.replace('.pdf', '_converted.docx')

    print("=" * 80)
    print("PDF TO DOCX CONVERSION (Smallpdf Style)")
    print("=" * 80)
    print(f"\nInput:  {pdf_path}")
    print(f"Output: {output_path}")

    # Step 1: Convert PDF to DOCX
    print("\n" + "=" * 80)
    print("STEP 1: Converting PDF to DOCX")
    print("=" * 80)

    converter = HybridConverter()
    converter.convert(pdf_path, output_path, use_yolo_images=True)

    # Step 2: Apply formatting
    if apply_formatting:
        print("\n" + "=" * 80)
        print("STEP 2: Applying Formatting")
        print("=" * 80)

        format_to_smallpdf_style(output_path)

    # Summary
    print("\n" + "=" * 80)
    print("CONVERSION COMPLETED!")
    print("=" * 80)
    print(f"\nOutput file: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / (1024*1024):.2f} MB")
    print("\nFormatting applied:")
    print("  - Font: Bookman Old Style, 10pt")
    print("  - Body text: JUSTIFY alignment")
    print("  - Sections: LEFT alignment")
    print("  - Figures: CENTER alignment")
    print("  - Math formulas: Preserved as-is")
    print("\n" + "=" * 80)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python convert_with_format.py <pdf_file> [output.docx]")
        print("\nExample:")
        print("  python convert_with_format.py input.pdf")
        print("  python convert_with_format.py input.pdf output.docx")
        sys.exit(1)

    pdf_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(pdf_file):
        print(f"Error: File not found: {pdf_file}")
        sys.exit(1)

    convert_pdf_to_docx(pdf_file, output_file)
