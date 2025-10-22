"""
Check all paragraphs around formulas in detail
"""
from docx import Document

docx_path = "test_files/result_hybrid.docx"

doc = Document(docx_path)

with open("check_all_paragraphs.txt", "w", encoding="utf-8") as f:
    f.write("All paragraphs from 30 to 45:\n")
    f.write("=" * 80 + "\n\n")

    for idx in range(30, min(50, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        text = para.text

        f.write(f"Para {idx}: {repr(text)}\n")

        # Check if contains 'groud' or 'velocity'
        if 'groud' in text.lower() or 'velocity' in text.lower():
            f.write(f"  [MATCH] Contains target keywords!\n")

    f.write("\n" + "=" * 80 + "\n")
    f.write("Also checking tables...\n\n")

    for table_idx, table in enumerate(doc.tables):
        if table_idx < 5:  # Check first 5 tables
            f.write(f"\nTable {table_idx}:\n")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        f.write(f"  Row {row_idx}, Cell {cell_idx}: {repr(text[:100])}\n")

print("Results written to check_all_paragraphs.txt")
