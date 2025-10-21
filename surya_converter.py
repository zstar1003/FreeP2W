"""
PDF to DOCX Converter using Surya Layout Detection

This converter uses AI-based layout detection (Surya) to:
1. Detect image regions accurately
2. Extract complete images without splitting
3. Parse text regions normally with pdf2docx
"""

import fitz
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from surya.foundation import FoundationPredictor
from surya.layout import LayoutPredictor
import io
import logging
from pathlib import Path

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')


class SuryaLayoutConverter:
    """PDF to DOCX converter using Surya AI layout detection"""

    def __init__(self):
        """Initialize Surya models"""
        logging.info("Loading Surya layout detection model...")
        # Initialize Foundation Predictor
        self.foundation_predictor = FoundationPredictor(device='cpu')  # Use 'cuda' if GPU available
        # Initialize Layout Predictor
        self.layout_predictor = LayoutPredictor(self.foundation_predictor)
        logging.info("Models loaded successfully")

    def convert(self, pdf_path, docx_path, dpi=150, image_quality=4.0):
        """
        Convert PDF to DOCX using layout detection

        Args:
            pdf_path (str): Path to input PDF file
            docx_path (str): Path to output DOCX file
            dpi (int): DPI for PDF to image conversion (default: 150)
            image_quality (float): Quality multiplier for extracted images (default: 4.0)
        """
        logging.info(f"Starting conversion: {pdf_path}")

        # Open PDF
        pdf_doc = fitz.open(pdf_path)
        docx_doc = Document()

        # Process each page
        for page_num in range(len(pdf_doc)):
            logging.info(f"Processing page {page_num + 1}/{len(pdf_doc)}")

            page = pdf_doc[page_num]

            # Step 1: Detect layout with Surya
            layout_regions = self._detect_layout(page, dpi)

            # Step 2: Sort regions by vertical position (reading order)
            layout_regions.sort(key=lambda r: (r['bbox'][1], r['bbox'][0]))

            # Step 3: Process each region
            for region in layout_regions:
                region_type = region['type']
                bbox = region['bbox']

                logging.info(f"  Region: {region_type} at {bbox}")

                if region_type in ['Picture', 'Figure']:
                    # Extract and insert image
                    self._insert_image_region(page, bbox, docx_doc, image_quality)

                elif region_type in ['Table']:
                    # Extract table (for now, as image)
                    self._insert_image_region(page, bbox, docx_doc, image_quality)

                elif region_type in ['Text', 'Title', 'Section-header', 'List-item', 'Caption']:
                    # Extract and insert text
                    self._insert_text_region(page, bbox, docx_doc, region_type)

                else:
                    logging.info(f"  Skipping region type: {region_type}")

            # Add page break (except for last page)
            if page_num < len(pdf_doc) - 1:
                docx_doc.add_page_break()

        # Save DOCX
        pdf_doc.close()
        docx_doc.save(docx_path)
        logging.info(f"Conversion completed: {docx_path}")

    def _detect_layout(self, page, dpi=150):
        """
        Detect layout regions using Surya

        Args:
            page (fitz.Page): PDF page
            dpi (int): DPI for rendering

        Returns:
            list: List of detected regions with type and bbox
        """
        # Convert PDF page to PIL Image
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        pil_image = Image.open(io.BytesIO(img_data))

        # Detect layout using Surya
        layout_results = self.layout_predictor([pil_image])

        # Convert results to our format
        regions = []
        page_width = page.rect.width
        page_height = page.rect.height

        for detection in layout_results[0].bboxes:
            # Surya returns pixel coordinates, convert to PDF coordinates
            x0, y0, x1, y1 = detection.bbox

            # Scale from image coordinates to PDF coordinates
            scale_x = page_width / pil_image.width
            scale_y = page_height / pil_image.height

            bbox = [
                x0 * scale_x,
                y0 * scale_y,
                x1 * scale_x,
                y1 * scale_y
            ]

            regions.append({
                'type': detection.label,
                'bbox': bbox,
                'confidence': detection.confidence if hasattr(detection, 'confidence') else 1.0
            })

        logging.info(f"  Detected {len(regions)} regions")
        return regions

    def _insert_image_region(self, page, bbox, docx_doc, quality=4.0):
        """
        Extract image region from PDF and insert into DOCX

        Args:
            page (fitz.Page): PDF page
            bbox (list): Bounding box [x0, y0, x1, y1]
            docx_doc (Document): DOCX document
            quality (float): Image quality multiplier
        """
        # Create rect from bbox
        rect = fitz.Rect(bbox)

        # Clip page region as image
        mat = fitz.Matrix(quality, quality)
        pix = page.get_pixmap(matrix=mat, clip=rect)

        # Convert to PNG bytes
        img_data = pix.tobytes("png")

        # Insert into DOCX
        try:
            img_stream = io.BytesIO(img_data)

            # Calculate width in inches (fit to page width if too large)
            max_width = 6.0  # inches
            width_pt = rect.width
            width_inches = width_pt / 72.0

            if width_inches > max_width:
                width_inches = max_width

            docx_doc.add_picture(img_stream, width=Inches(width_inches))
            logging.info(f"    Inserted image: {rect.width:.1f} x {rect.height:.1f} pt")

        except Exception as e:
            logging.error(f"    Failed to insert image: {e}")

    def _insert_text_region(self, page, bbox, docx_doc, region_type):
        """
        Extract text region from PDF and insert into DOCX

        Args:
            page (fitz.Page): PDF page
            bbox (list): Bounding box [x0, y0, x1, y1]
            docx_doc (Document): DOCX document
            region_type (str): Type of text region
        """
        # Extract text from region
        rect = fitz.Rect(bbox)
        text = page.get_text("text", clip=rect).strip()

        if not text:
            return

        # Add to DOCX with appropriate styling
        if region_type in ['Title', 'Section-header']:
            para = docx_doc.add_heading(text, level=1 if region_type == 'Title' else 2)
        else:
            para = docx_doc.add_paragraph(text)

        logging.info(f"    Inserted text: {len(text)} chars")


def main():
    """Test the converter"""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python surya_converter.py <input.pdf> [output.docx]")
        print("\nExample:")
        print("  python surya_converter.py test_files/2503.20314v2.pdf test_files/result_surya.docx")
        return

    pdf_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) > 2 else pdf_path.replace('.pdf', '_surya.docx')

    converter = SuryaLayoutConverter()
    converter.convert(pdf_path, docx_path)

    print(f"\nConversion complete!")
    print(f"Output: {docx_path}")


if __name__ == "__main__":
    main()
