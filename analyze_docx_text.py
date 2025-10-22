"""
Analyze text content in DOCX file
"""
from docx import Document

def analyze_docx_text(docx_path):
    """Analyze text content in a DOCX file"""
    doc = Document(docx_path)

    total_paragraphs = len(doc.paragraphs)
    non_empty_paragraphs = 0
    total_chars = 0

    print(f"Analyzing {docx_path}...")
    print(f"\nTotal paragraphs: {total_paragraphs}")

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            non_empty_paragraphs += 1
            total_chars += len(text)
            if i < 10:  # Show first 10 non-empty paragraphs
                print(f"  Para {i+1}: {text[:100]}{'...' if len(text) > 100 else ''}")

    print(f"\nNon-empty paragraphs: {non_empty_paragraphs}")
    print(f"Total characters: {total_chars}")
    print(f"Total inline shapes (images): {len(doc.inline_shapes)}")

    return non_empty_paragraphs, total_chars

if __name__ == "__main__":
    analyze_docx_text("test_files/result_hybrid.docx")
