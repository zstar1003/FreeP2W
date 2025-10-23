"""
Check if Unicode math symbols are preserved in the DOCX
"""
from docx import Document

docx_path = "test_files/result_hybrid.docx"

doc = Document(docx_path)

print("Searching for paragraphs with math symbols...")
print("=" * 80)

# Keywords to search for
math_keywords = ['latent space', 'velocity', 'model weights', 'embedding']

for idx, para in enumerate(doc.paragraphs):
    text = para.text

    # Check if contains any math keywords
    for keyword in math_keywords:
        if keyword in text.lower():
            print(f"\nPara {idx}: {keyword}")
            print(f"Text: {text[:200]}")

            # Check for Unicode math symbols
            has_unicode = False
            for char in text:
                if ord(char) > 127:  # Non-ASCII character
                    print(f"  Unicode char: {char} (U+{ord(char):04X})")
                    has_unicode = True

            if not has_unicode:
                print("  No Unicode symbols detected")

            print("-" * 80)

print("\n" + "=" * 80)
print("Checking for replacement character �")

for idx, para in enumerate(doc.paragraphs):
    if '�' in para.text or '\ufffd' in para.text:
        print(f"\nPara {idx} contains �:")
        print(f"  {para.text[:200]}")
