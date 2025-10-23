"""
Compare our minimal formatted output with smallpdf reference
Focus on key differences that matter
"""

from docx import Document
from collections import Counter
import re

def deep_analyze_doc(path, doc_name):
    """Deep analysis of document structure"""
    doc = Document(path)

    analysis = {
        'total_paragraphs': len(doc.paragraphs),
        'alignments': [],
        'font_names': [],
        'font_sizes': [],
        'none_fonts': 0,
        'set_fonts': 0,
        'spacing_info': [],
        'indent_info': []
    }

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Alignment
        if para.alignment is not None:
            analysis['alignments'].append(para.alignment)
        else:
            analysis['alignments'].append('None')

        # Spacing
        if para.paragraph_format.space_before:
            analysis['spacing_info'].append('has_space_before')
        if para.paragraph_format.space_after:
            analysis['spacing_info'].append('has_space_after')
        if para.paragraph_format.line_spacing:
            analysis['spacing_info'].append('has_line_spacing')

        # Indents
        if para.paragraph_format.first_line_indent:
            analysis['indent_info'].append('has_first_line')
        if para.paragraph_format.left_indent:
            analysis['indent_info'].append('has_left')
        if para.paragraph_format.right_indent:
            analysis['indent_info'].append('has_right')

        # Fonts
        for run in para.runs:
            if run.font.name is None or run.font.name == '':
                analysis['none_fonts'] += 1
            else:
                analysis['set_fonts'] += 1
                analysis['font_names'].append(run.font.name)

            if run.font.size:
                analysis['font_sizes'].append(run.font.size.pt)

    print(f"\n{'=' * 80}")
    print(f"{doc_name.upper()}")
    print(f"{'=' * 80}")
    print(f"Total paragraphs: {analysis['total_paragraphs']}")
    print(f"\nFont information:")
    print(f"  None/empty fonts: {analysis['none_fonts']}")
    print(f"  Set fonts: {analysis['set_fonts']}")

    if analysis['font_names']:
        font_counter = Counter(analysis['font_names'])
        print(f"  Font name distribution:")
        for font, count in font_counter.most_common(10):
            pct = count / analysis['set_fonts'] * 100
            print(f"    {font}: {count} ({pct:.1f}%)")

    if analysis['font_sizes']:
        size_counter = Counter(analysis['font_sizes'])
        print(f"\n  Font size distribution:")
        for size, count in size_counter.most_common(10):
            print(f"    {size}pt: {count}")

    align_counter = Counter(analysis['alignments'])
    print(f"\nAlignment distribution:")
    for align, count in align_counter.most_common():
        print(f"  {align}: {count}")

    spacing_counter = Counter(analysis['spacing_info'])
    if spacing_counter:
        print(f"\nSpacing properties:")
        for prop, count in spacing_counter.most_common():
            print(f"  {prop}: {count} paragraphs")

    indent_counter = Counter(analysis['indent_info'])
    if indent_counter:
        print(f"\nIndent properties:")
        for prop, count in indent_counter.most_common():
            print(f"  {prop}: {count} paragraphs")

    return analysis

print("=" * 80)
print("COMPARING OUR OUTPUT WITH SMALLPDF REFERENCE")
print("=" * 80)

# Analyze both documents
smallpdf = deep_analyze_doc('test_files/2503.20314v2_smallpdf.docx', 'Smallpdf Reference')
ours = deep_analyze_doc('result_minimal_format.docx', 'Our Minimal Format Output')

# Key comparisons
print("\n" + "=" * 80)
print("KEY DIFFERENCES")
print("=" * 80)

print("\n1. FONT NAMES:")
print("-" * 80)
smallpdf_fonts = set([f for f in smallpdf['font_names'] if f])
our_fonts = set([f for f in ours['font_names'] if f])

print(f"Smallpdf has {len(smallpdf_fonts)} unique fonts: {sorted(smallpdf_fonts)}")
print(f"Ours has {len(our_fonts)} unique fonts: {sorted(our_fonts)}")

if our_fonts - smallpdf_fonts:
    print(f"\nFonts we have but smallpdf doesn't: {our_fonts - smallpdf_fonts}")
if smallpdf_fonts - our_fonts:
    print(f"Fonts smallpdf has but we don't: {smallpdf_fonts - our_fonts}")

print("\n2. NONE/EMPTY FONTS:")
print("-" * 80)
smallpdf_none_pct = smallpdf['none_fonts'] / (smallpdf['none_fonts'] + smallpdf['set_fonts']) * 100
ours_none_pct = ours['none_fonts'] / (ours['none_fonts'] + ours['set_fonts']) * 100

print(f"Smallpdf: {smallpdf['none_fonts']}/{smallpdf['none_fonts'] + smallpdf['set_fonts']} ({smallpdf_none_pct:.1f}%) None/empty")
print(f"Ours: {ours['none_fonts']}/{ours['none_fonts'] + ours['set_fonts']} ({ours_none_pct:.1f}%) None/empty")

if abs(smallpdf_none_pct - ours_none_pct) > 10:
    print("\nWARNING: Large difference in None font percentage!")
    print("This suggests we may be setting fonts where smallpdf doesn't, or vice versa")

print("\n3. ALIGNMENT:")
print("-" * 80)
smallpdf_align = Counter(smallpdf['alignments'])
ours_align = Counter(ours['alignments'])

print("Smallpdf alignment:")
for align, count in smallpdf_align.most_common():
    pct = count / len(smallpdf['alignments']) * 100
    print(f"  {align}: {count} ({pct:.1f}%)")

print("\nOurs alignment:")
for align, count in ours_align.most_common():
    pct = count / len(ours['alignments']) * 100
    print(f"  {align}: {count} ({pct:.1f}%)")

print("\n" + "=" * 80)
print("CONCLUSION")
print("=" * 80)

# Calculate overall similarity
issues = []

if abs(smallpdf_none_pct - ours_none_pct) > 10:
    issues.append("- Large difference in None font percentage")

if len(smallpdf_fonts - our_fonts) > 2:
    issues.append(f"- Missing {len(smallpdf_fonts - our_fonts)} fonts that smallpdf has")

if len(our_fonts - smallpdf_fonts) > 2:
    issues.append(f"- Have {len(our_fonts - smallpdf_fonts)} extra fonts that smallpdf doesn't have")

align_diff = abs(smallpdf_align.get(3, 0) - ours_align.get(3, 0))  # JUSTIFY alignment
if align_diff > 5:
    issues.append(f"- Different JUSTIFY alignment count (diff: {align_diff})")

if issues:
    print("\nRemaining issues:")
    for issue in issues:
        print(issue)
else:
    print("\nNo major differences found! Our formatting appears to match smallpdf well.")

print("=" * 80)
