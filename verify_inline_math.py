"""
Verify inline math symbols in the processed DOCX
"""
from docx import Document
from docx.oxml.ns import qn

docx_path = "test_files/result_hybrid_with_inline_math.docx"

doc = Document(docx_path)

print("Checking for inline math in processed DOCX...")
print("=" * 80)

# Keywords to search for areas with math
math_keywords = ['latent space', 'velocity', 'model weights', 'embedding']

for idx, para in enumerate(doc.paragraphs):
    text = para.text

    # Check if contains math keywords
    for keyword in math_keywords:
        if keyword in text.lower():
            print(f"\nPara {idx}: {text[:150]}...")

            # Check for math elements in XML
            para_elem = para._element
            math_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')

            if math_elems:
                print(f"  ✓ Contains {len(math_elems)} inline math element(s)")

                # Show first math element content
                for math_elem in math_elems[:3]:
                    math_text = ''.join(math_elem.itertext())
                    print(f"    Math content: {repr(math_text)}")
            else:
                print(f"  ✗ No inline math elements found")

            break

print("\n" + "=" * 80)
print("Summary: Checking specific paragraph with known math symbols...")

# Check paragraph 23 which should have x ∈ R...
for idx in [23, 26, 41]:
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text
        para_elem = para._element
        math_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')

        print(f"\nPara {idx}:")
        print(f"  Text: {text[:100]}...")
        print(f"  Math elements: {len(math_elems)}")
