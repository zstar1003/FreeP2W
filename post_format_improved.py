"""
改进的后处理格式化脚本 - 确保对齐方式全覆盖

改进：
1. 对正文段落：设置字体 + 大小 + 对齐
2. 对非正文但包含英文的段落：至少设置对齐方式
3. 只有明确的标题、目录才跳过对齐设置
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


def is_title_or_heading(para):
    """判断段落是否为标题"""
    text = para.text.strip()
    if not text:
        return True

    # 全大写短文本
    if len(text) < 100 and text.isupper():
        return True

    # 检查字体大小和粗体
    for run in para.runs:
        if run.font.size and run.font.size.pt >= 14:
            return True
        if run.font.bold and len(text) < 100:
            return True

    return False


def is_figure_or_table_caption(para):
    """判断是否为图表标题"""
    text = para.text.strip()
    patterns = [
        r'^Figure\s+\d+',
        r'^Fig\.\s*\d+',
        r'^Table\s+\d+',
        r'^图\s*\d+',
        r'^表\s*\d+',
    ]
    for pattern in patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    return False


def detect_toc_range(doc):
    """检测目录范围"""
    toc_start = None
    toc_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if toc_start is None:
            if 'contents' in text or 'table of contents' in text or '目录' in text:
                toc_start = i
                continue

        if toc_start is not None and toc_end is None:
            if len(para.text) > 100 or any(keyword in text for keyword in ['abstract', 'introduction', '摘要', '引言']):
                toc_end = i
                break

    if toc_start is not None and toc_end is None:
        toc_end = min(toc_start + 20, len(doc.paragraphs))

    if toc_start is not None:
        return (toc_start, toc_end)

    return None


def is_toc_paragraph(para, para_idx, toc_range):
    """判断是否在目录范围内"""
    if toc_range is None:
        return False
    start, end = toc_range
    return start <= para_idx < end


def has_math_content(para):
    """判断段落是否包含数学公式（OMML）"""
    return 'oMath' in para._element.xml if hasattr(para, '_element') else False


def contains_english(text):
    """判断文本是否包含英文"""
    return bool(re.search(r'[a-zA-Z]', text))


def should_skip_completely(para, para_idx, toc_range):
    """
    判断是否应该完全跳过（连对齐都不设置）

    只有以下情况才完全跳过：
    - 标题（保持原格式）
    - 目录（保持左对齐）
    """
    # 跳过标题
    if is_title_or_heading(para):
        return True

    # 跳过目录
    if is_toc_paragraph(para, para_idx, toc_range):
        return True

    return False


def is_body_paragraph_strict(para, para_idx, toc_range):
    """
    严格判断是否为正文段落（需要全面格式化）
    """
    text = para.text.strip()

    if not text or len(text) < 20:
        return False

    if is_title_or_heading(para):
        return False

    if is_figure_or_table_caption(para):
        return False

    if is_toc_paragraph(para, para_idx, toc_range):
        return False

    return True


def format_english_body_text_improved(input_path, output_path,
                                       target_font='Times New Roman',
                                       target_size=10.0,
                                       target_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """
    改进的格式化函数

    策略：
    1. 正文段落：字体 + 大小 + 对齐（全格式化）
    2. 其他英文段落：至少设置对齐方式
    3. 只有标题和目录完全不动
    """
    print("=" * 80)
    print("改进的后处理格式化 - 全覆盖对齐方式")
    print("=" * 80)
    print(f"\n输入文件: {input_path}")
    print(f"输出文件: {output_path}")
    print(f"\n目标格式:")
    print(f"  字体: {target_font}")
    print(f"  大小: {target_size}pt")
    print(f"  对齐: JUSTIFY (两端对齐)")
    print("\n" + "=" * 80)

    doc = Document(input_path)

    # 检测目录范围
    toc_range = detect_toc_range(doc)
    if toc_range:
        print(f"\n[检测] 目录范围: 段落 {toc_range[0]} - {toc_range[1]}")
    else:
        print(f"\n[检测] 未检测到目录")

    # 统计信息
    stats = {
        'total_paragraphs': len(doc.paragraphs),
        'full_formatted': 0,        # 完全格式化（字体+大小+对齐）
        'alignment_only': 0,         # 只设置对齐
        'skipped_title': 0,
        'skipped_toc': 0,
        'skipped_math': 0,
        'skipped_no_english': 0,
        'formatted_runs': 0,
    }

    print(f"\n[处理] 开始格式化 {stats['total_paragraphs']} 个段落...")
    print()

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 跳过空段落
        if not text:
            continue

        # 检查是否包含英文
        if not contains_english(text):
            stats['skipped_no_english'] += 1
            continue

        # 检查是否应该完全跳过
        if should_skip_completely(para, para_idx, toc_range):
            if is_title_or_heading(para):
                stats['skipped_title'] += 1
            elif is_toc_paragraph(para, para_idx, toc_range):
                stats['skipped_toc'] += 1
            continue

        # 判断是否包含数学公式
        has_math = has_math_content(para)

        # 判断是否为严格意义的正文
        is_strict_body = is_body_paragraph_strict(para, para_idx, toc_range)

        if is_strict_body and not has_math:
            # 正文且无公式：完全格式化
            para.alignment = target_alignment

            for run in para.runs:
                if contains_english(run.text):
                    run.font.name = target_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), target_font)
                    run.font.size = Pt(target_size)
                    stats['formatted_runs'] += 1

            stats['full_formatted'] += 1

        else:
            # 非典型正文 或 包含公式：至少设置对齐
            # （如图表标题、过短段落、包含公式的段落等）
            para.alignment = target_alignment

            # 如果包含公式，记录
            if has_math:
                stats['skipped_math'] += 1

            stats['alignment_only'] += 1

        # 进度显示
        if (stats['full_formatted'] + stats['alignment_only']) % 10 == 0:
            print(f"  已处理 {stats['full_formatted'] + stats['alignment_only']} 个段落...")

    print(f"\n[完成] 格式化统计:")
    print("=" * 80)
    print(f"总段落数:          {stats['total_paragraphs']}")
    print(f"完全格式化:        {stats['full_formatted']}  (字体+大小+对齐)")
    print(f"仅设置对齐:        {stats['alignment_only']}  (保留原字体和大小)")
    print(f"已格式化 runs:     {stats['formatted_runs']}")
    print()
    print(f"跳过的段落:")
    print(f"  - 标题:          {stats['skipped_title']}")
    print(f"  - 目录:          {stats['skipped_toc']}")
    print(f"  - 包含公式:      {stats['skipped_math']}")
    print(f"  - 无英文:        {stats['skipped_no_english']}")
    print("=" * 80)

    # 保存文档
    doc.save(output_path)
    print(f"\n[保存] 已保存到: {output_path}")
    print("=" * 80)


if __name__ == "__main__":
    import sys

    input_file = "result.docx"
    output_file = "result_formatted.docx"

    # 支持命令行参数
    if len(sys.argv) >= 2:
        input_file = sys.argv[1]
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]

    format_english_body_text_improved(
        input_path=input_file,
        output_path=output_file,
        target_font='Times New Roman',
        target_size=10.0,
        target_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    print("\n后续步骤:")
    print(f"  1. 打开 {output_file} 检查效果")
    print("  2. 运行 python diagnose_formatting.py 诊断问题")
    print("  3. 运行 python test_formatting_effect.py 查看对比")
