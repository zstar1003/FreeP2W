"""
Format DOCX to match smallpdf style
Key characteristics from smallpdf analysis:
1. Font: Bookman Old Style (primary), 10pt
2. Alignment: mostly None (left-aligned, not justified)
3. Line spacing: various values
4. Section headers: None alignment, 10pt
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
import re

def format_to_smallpdf_style(input_path, output_path):
    """
    Format DOCX to match smallpdf.docx style:
    - Use Bookman Old Style font
    - Set 10pt as standard body size
    - Use LEFT (None) alignment for body text (NOT justify)
    - Section headers: LEFT, 10pt
    - Figure captions: CENTER (keep this)
    - Clean TOC spacing
    """

    doc = Document(input_path)

    print("=" * 80)
    print("FORMATTING TO SMALLPDF STYLE")
    print("=" * 80)

    # Determine body text size
    font_sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                font_sizes.append(run.font.size.pt)

    body_size = Counter(font_sizes).most_common(1)[0][0] if font_sizes else 10.0
    print(f"\nBody text size: {body_size}pt")

    # Detect TOC
    toc_start = None
    toc_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if toc_start is None and ('contents' in text or 'table of contents' in text):
            toc_start = i
            print(f"TOC starts at paragraph {i}")

        if toc_start is not None and toc_end is None and i > toc_start:
            para_text = para.text.strip()
            if (re.match(r'^\d+\.?\d*\s+[A-Z]', para_text) or
                any(keyword in text for keyword in ['abstract', 'introduction']) or
                len(para_text) > 100):
                toc_end = i
                print(f"TOC ends at paragraph {i}")
                break

    if toc_start and not toc_end:
        toc_end = min(toc_start + 15, len(doc.paragraphs))

    # Format paragraphs
    stats = {
        'titles': 0,
        'sections': 0,
        'figures': 0,
        'body': 0,
        'toc': 0,
        'math': 0
    }

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check if paragraph contains math
        has_math = 'oMath' in para._element.xml if hasattr(para, '_element') else False
        if has_math:
            stats['math'] += 1
            continue  # Don't modify math paragraphs

        # Detect paragraph type
        max_font_size = 0
        for run in para.runs:
            if run.font.size:
                max_font_size = max(max_font_size, run.font.size.pt)

        is_figure = text.lower().startswith(('figure', 'fig.', 'fig ', 'table'))
        is_section = bool(re.match(r'^[\d\.]+\s+[A-Z]', text))
        is_title = max_font_size > 12 or (len(text) < 100 and text.isupper() and not is_figure)
        in_toc = toc_start is not None and toc_start <= i < toc_end

        # Apply formatting based on smallpdf style
        if in_toc:
            # TOC: LEFT align, clean spacing, body font
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cleaned = re.sub(r'\s+', ' ', text).strip()
            if cleaned != text and para.runs:
                para.clear()
                run = para.add_run(cleaned)
                run.font.name = 'Bookman Old Style'
                run.font.size = Pt(body_size)
            stats['toc'] += 1

        elif is_figure:
            # Figure captions: CENTER align (keep this different from smallpdf)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Bookman Old Style'
                if run.font.size and run.font.size.pt != body_size:
                    run.font.size = Pt(body_size)
            stats['figures'] += 1

        elif is_section:
            # Section headers: LEFT align (None), keep size or set to body
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = 'Bookman Old Style'
                # Section headers in smallpdf are 10pt
                if run.font.size and run.font.size.pt <= 12:
                    run.font.size = Pt(body_size)
            stats['sections'] += 1

        elif is_title:
            # Main titles: keep original size and LEFT align
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Don't change font for titles (keep original)
            stats['titles'] += 1

        else:
            # Body text: use JUSTIFY align (matches smallpdf style)
            # Smallpdf uses JUSTIFY for most body paragraphs
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Set font and size
            for run in para.runs:
                run.font.name = 'Bookman Old Style'
                if run.font.size and run.font.size.pt <= 12:
                    run.font.size = Pt(body_size)

            stats['body'] += 1

    print(f"\nFormatting summary:")
    print(f"  - Titles: {stats['titles']}")
    print(f"  - Section headers: {stats['sections']}")
    print(f"  - Figure captions: {stats['figures']}")
    print(f"  - Body paragraphs: {stats['body']}")
    print(f"  - TOC entries: {stats['toc']}")
    print(f"  - Math formulas: {stats['math']} (skipped)")

    # Save
    doc.save(output_path)
    print(f"\nSaved to: {output_path}")
    print("=" * 80)

if __name__ == "__main__":
    format_to_smallpdf_style("result.docx", "result_smallpdf_style.docx")
