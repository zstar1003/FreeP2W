"""
Check if the missing text exists in the DOCX output
"""
from docx import Document

docx_path = "test_files/result_hybrid.docx"

doc = Document(docx_path)

print("Searching for 'velocity' or 'ground truth' in DOCX...\n")
print("=" * 80)

found_count = 0

for idx, para in enumerate(doc.paragraphs):
    text = para.text
    if 'velocity' in text.lower() or 'ground truth' in text.lower():
        found_count += 1
        print(f"\nParagraph {idx}:")
        print(f"  Text: {text[:200]}")

# Also check tables
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text
            if 'velocity' in text.lower() or 'ground truth' in text.lower():
                found_count += 1
                print(f"\nTable {table_idx}, Row {row_idx}, Cell {cell_idx}:")
                print(f"  Text: {text[:200]}")

print("\n" + "=" * 80)
print(f"Total occurrences found: {found_count}")

if found_count == 0:
    print("\n[WARNING] Text containing 'velocity' or 'ground truth' NOT FOUND in DOCX!")
    print("This confirms the text is missing from the output.")

print("=" * 80)
