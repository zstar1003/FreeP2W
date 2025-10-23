"""
Find paragraphs with math symbols in original vs processed DOCX
"""
from docx import Document

print("=" * 80)
print("ORIGINAL FILE (result_hybrid.docx)")
print("=" * 80)

doc_orig = Document("test_files/result_hybrid.docx")

for idx, para in enumerate(doc_orig.paragraphs):
    text = para.text
    if 'latent space' in text.lower():
        print(f"\nPara {idx}: {text[:200]}")
        # Check for special characters
        has_unicode = any(ord(c) > 127 for c in text)
        if has_unicode:
            print("  Contains Unicode symbols")

print("\n" + "=" * 80)
print("PROCESSED FILE (result_hybrid_with_inline_math.docx)")
print("=" * 80)

doc_proc = Document("test_files/result_hybrid_with_inline_math.docx")

for idx, para in enumerate(doc_proc.paragraphs):
    text = para.text
    if 'latent space' in text.lower():
        print(f"\nPara {idx}: {text[:200]}")

        # Check for math elements
        para_elem = para._element
        math_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
        print(f"  Math elements: {len(math_elems)}")

        # Count runs
        runs = para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        print(f"  Total runs: {len(runs)}")
