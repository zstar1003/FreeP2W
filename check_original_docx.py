"""
Check if the text exists in pure pdf2docx output (without hybrid method)
"""
from docx import Document

docx_path = "test_files/result_pdf2docx_origin.docx"

doc = Document(docx_path)

print("Searching for 'groud truth velocity' in pure pdf2docx output...\n")
print("=" * 80)

found_count = 0

for idx, para in enumerate(doc.paragraphs):
    text = para.text
    if 'groud truth' in text.lower() or ('velocity' in text.lower() and 'vt is' in text.lower()):
        found_count += 1
        print(f"\nParagraph {idx}:")
        print(f"  Text: {text[:300].encode('utf-8', errors='replace').decode('utf-8')}")

# Also check context around formula paragraphs
print("\n\nParagraphs around formulas (35-40):")
print("=" * 80)

for idx in range(30, min(45, len(doc.paragraphs))):
    para = doc.paragraphs[idx]
    text = para.text.strip()
    if text:
        print(f"\nPara {idx}: {text[:150]}")

print("\n" + "=" * 80)
print(f"Total 'groud truth velocity' occurrences: {found_count}")

if found_count == 0:
    print("\n[FINDING] Text 'The groud truth velocity vt is' is ALSO MISSING in pure pdf2docx output!")
    print("This suggests the problem is with pdf2docx, NOT our hybrid method.")
else:
    print("\n[FINDING] Text exists in pure pdf2docx but missing in hybrid output!")
    print("This suggests our hybrid method is causing the problem.")

print("=" * 80)
