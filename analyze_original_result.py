"""
Analyze our original result.docx to understand the actual format
"""

from docx import Document
from docx.shared import Pt

doc = Document('result.docx')

print("="*80)
print("ORIGINAL RESULT.DOCX ANALYSIS")
print("="*80)

with open('original_result_analysis.txt', 'w', encoding='utf-8') as f:
    f.write("="*80 + "\n")
    f.write("ORIGINAL RESULT.DOCX DETAILED ANALYSIS\n")
    f.write("="*80 + "\n\n")

    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text.strip()
        if not text or len(text) < 20:
            continue

        f.write(f"\n{'='*80}\n")
        f.write(f"PARAGRAPH {i}: {text[:70]}...\n")
        f.write(f"{'='*80}\n")

        # Paragraph alignment
        alignment = str(para.alignment).split('.')[-1] if para.alignment else 'None'
        f.write(f"Alignment: {alignment}\n")

        # Paragraph spacing
        space_before = para.paragraph_format.space_before
        space_after = para.paragraph_format.space_after
        line_spacing = para.paragraph_format.line_spacing

        f.write(f"Space before: {space_before}\n")
        f.write(f"Space after: {space_after}\n")
        f.write(f"Line spacing: {line_spacing}\n")

        # Check if has math
        has_math = 'oMath' in para._element.xml
        f.write(f"Has math: {has_math}\n")

        # Run details
        f.write(f"\nRuns ({len(para.runs)} total):\n")
        for j, run in enumerate(para.runs[:5]):  # First 5 runs
            font_name = run.font.name if run.font.name else 'None'
            font_size = f"{run.font.size.pt}pt" if run.font.size else 'None'
            run_text = run.text[:30]
            f.write(f"  Run {j}: font={font_name:15s} size={font_size:8s} text='{run_text}'\n")

print("Analysis saved to: original_result_analysis.txt")

# Quick summary
print("\nQUICK SUMMARY:")
print("-"*80)

alignment_stats = {}
font_stats = {}
size_stats = {}

for para in doc.paragraphs[:50]:
    text = para.text.strip()
    if not text:
        continue

    # Alignment
    if para.alignment:
        alignment_stats[str(para.alignment)] = alignment_stats.get(str(para.alignment), 0) + 1
    else:
        alignment_stats['None'] = alignment_stats.get('None', 0) + 1

    # Fonts
    for run in para.runs:
        if run.font.name:
            font_stats[run.font.name] = font_stats.get(run.font.name, 0) + 1
        else:
            font_stats['None'] = font_stats.get('None', 0) + 1

        if run.font.size:
            size_stats[run.font.size.pt] = size_stats.get(run.font.size.pt, 0) + 1

print("\nAlignment distribution:")
for align, count in sorted(alignment_stats.items(), key=lambda x: -x[1])[:5]:
    print(f"  {align:20s}: {count}")

print("\nFont name distribution:")
for font, count in sorted(font_stats.items(), key=lambda x: -x[1])[:5]:
    print(f"  {font:20s}: {count}")

print("\nFont size distribution:")
for size, count in sorted(size_stats.items(), key=lambda x: -x[1] if x[0] != 'None' else 0)[:5]:
    print(f"  {size}pt: {count}")
