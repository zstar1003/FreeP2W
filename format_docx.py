"""
Format DOCX file: unify fonts, fix alignment, clean up TOC
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
import re

def format_docx(input_path, output_path):
    """
    Format DOCX file:
    1. Unify body text to 10pt (most common size)
    2. Set body paragraphs to JUSTIFY alignment
    3. Detect and format TOC section (LEFT alignment, clean spacing)
    4. Keep title formatting (larger fonts)
    """

    doc = Document(input_path)

    print("=" * 80)
    print("FORMATTING DOCX FILE")
    print("=" * 80)

    # Analyze font sizes to determine body text size
    font_sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                font_sizes.append(run.font.size.pt)

    if font_sizes:
        body_size = Counter(font_sizes).most_common(1)[0][0]
        print(f"\nBody text size: {body_size}pt")
    else:
        body_size = 10.0
        print(f"\nUsing default body text size: {body_size}pt")

    # Detect TOC section
    toc_start = None
    toc_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        # Find "Contents" or "Table of Contents"
        if toc_start is None and ('contents' in text or 'table of contents' in text):
            toc_start = i
            print(f"\nTOC starts at paragraph {i}: '{para.text.strip()}'")

        # TOC usually ends before first section (4.x) or "Abstract" or "Introduction"
        if toc_start is not None and toc_end is None and i > toc_start:
            # Check if this looks like a section header or main content start
            para_text = para.text.strip()

            # End TOC if we see:
            # 1. Section numbers like "4.2 MODEL TRAINING"
            # 2. Keywords like "abstract", "introduction"
            # 3. A paragraph longer than 100 chars (likely body text)
            if (re.match(r'^\d+\.?\d*\s+[A-Z]', para_text) or
                any(keyword in text for keyword in ['abstract', 'introduction']) or
                len(para_text) > 100):
                toc_end = i
                print(f"TOC ends at paragraph {i}: '{para_text[:60]}...'")
                break

    if toc_start and not toc_end:
        toc_end = min(toc_start + 15, len(doc.paragraphs))
        print(f"TOC ends at paragraph {toc_end} (default)")

    # Format paragraphs
    title_count = 0
    body_count = 0
    toc_count = 0
    math_count = 0
    figure_count = 0
    section_count = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check if paragraph contains math
        has_math = 'oMath' in para._element.xml if hasattr(para, '_element') else False

        # Skip paragraphs with math formulas
        if has_math:
            math_count += 1
            continue

        # Detect paragraph type
        is_title = False
        is_figure = False
        is_section = False
        max_font_size = 0

        for run in para.runs:
            if run.font.size:
                size_pt = run.font.size.pt
                max_font_size = max(max_font_size, size_pt)

        # Check for figure captions (Figure X:)
        if text.lower().startswith('figure ') or text.lower().startswith('fig.') or text.lower().startswith('fig '):
            is_figure = True

        # Check for section numbers (4.2.1, 1.1, etc.) - more flexible pattern
        # Match: digits and dots, followed by space and text
        section_pattern = re.match(r'^[\d\.]+\s+[A-Z\-]', text)
        if section_pattern:
            is_section = True

        # Also match section headers like "4.2 MODEL TRAINING" (all caps with number)
        if re.match(r'^[\d\.]+\s+[A-Z][A-Z\s\-]+$', text):
            is_section = True

        # Consider as title if font > 12pt or all caps short text
        if max_font_size > 12 or (len(text) < 100 and text.isupper() and not is_figure):
            is_title = True

        # Format based on section and type
        if toc_start is not None and toc_start <= i < toc_end:
            # TOC section: left align, clean spacing
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Clean up extra spaces in TOC entries
            cleaned_text = re.sub(r'\s+', ' ', text)  # Multiple spaces -> single space
            cleaned_text = re.sub(r'^\s+', '', cleaned_text)  # Remove leading spaces
            cleaned_text = re.sub(r'\s+$', '', cleaned_text)  # Remove trailing spaces

            if cleaned_text != text and para.runs:
                # Rebuild paragraph with cleaned text
                para.clear()
                run = para.add_run(cleaned_text)
                run.font.size = Pt(body_size)

            toc_count += 1

        elif is_figure:
            # Figure captions: center align, keep or set to body size
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in para.runs:
                if run.font.size and run.font.size.pt != body_size:
                    run.font.size = Pt(body_size)

            figure_count += 1

        elif is_section:
            # Section titles: left align, keep current size
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_count += 1

        elif is_title:
            # Main titles: keep original formatting
            if not para.alignment or para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_count += 1

        else:
            # Body text: justify alignment, unify font size
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Unify font size for all runs
            for run in para.runs:
                if run.font.size and run.font.size.pt != body_size:
                    # Only change if not a title-sized font
                    if run.font.size.pt <= 12:
                        run.font.size = Pt(body_size)

            body_count += 1

    print(f"\nFormatting summary:")
    print(f"  - Titles: {title_count}")
    print(f"  - Section headers: {section_count}")
    print(f"  - Figure captions: {figure_count}")
    print(f"  - Body paragraphs: {body_count}")
    print(f"  - TOC entries: {toc_count}")
    print(f"  - Math formulas: {math_count} (skipped)")

    # Save
    doc.save(output_path)
    print(f"\nSaved to: {output_path}")
    print("=" * 80)

if __name__ == "__main__":
    format_docx("result.docx", "result_formatted.docx")
