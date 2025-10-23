"""
改进的公式插入方法 - 保留原始 run 结构

核心改进：
1. 不使用 para.clear() 破坏所有 runs
2. 精确定位包含占位符的 run
3. 只替换特定 run，保留其他 runs 的格式
4. 保持段落的原始 run 数量和格式属性
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import re


def find_run_with_placeholder(para, placeholder_id):
    """
    在段落中查找包含占位符的 run

    Args:
        para: docx.paragraph.Paragraph 对象
        placeholder_id: 占位符 ID 字符串

    Returns:
        tuple: (run_index, run, start_pos_in_run) 或 (None, None, None)
    """
    for run_idx, run in enumerate(para.runs):
        run_text = run.text
        if placeholder_id in run_text:
            start_pos = run_text.find(placeholder_id)
            return run_idx, run, start_pos

    return None, None, None


def preserve_run_formatting(source_run):
    """
    提取 run 的格式属性

    Args:
        source_run: 源 run 对象

    Returns:
        dict: 格式属性字典
    """
    formatting = {
        'font_name': source_run.font.name,
        'font_size': source_run.font.size,
        'bold': source_run.font.bold,
        'italic': source_run.font.italic,
        'underline': source_run.font.underline,
        'color': source_run.font.color.rgb if source_run.font.color.rgb else None,
    }
    return formatting


def apply_run_formatting(target_run, formatting):
    """
    应用格式到目标 run

    Args:
        target_run: 目标 run 对象
        formatting: 格式属性字典
    """
    if formatting['font_name']:
        target_run.font.name = formatting['font_name']
    if formatting['font_size']:
        target_run.font.size = formatting['font_size']
    if formatting['bold'] is not None:
        target_run.font.bold = formatting['bold']
    if formatting['italic'] is not None:
        target_run.font.italic = formatting['italic']
    if formatting['underline'] is not None:
        target_run.font.underline = formatting['underline']
    if formatting['color']:
        target_run.font.color.rgb = formatting['color']


def replace_placeholder_in_run(para, run_idx, run, start_pos, placeholder_id,
                                math_run_element, equation_number=None):
    """
    在特定 run 中替换占位符，保留其他 runs

    Args:
        para: 段落对象
        run_idx: run 索引
        run: 包含占位符的 run
        start_pos: 占位符在 run 中的起始位置
        placeholder_id: 占位符 ID
        math_run_element: OMML math run 元素
        equation_number: 公式编号（可选）

    Returns:
        bool: 是否成功替换
    """
    try:
        run_text = run.text
        text_before = run_text[:start_pos]
        text_after = run_text[start_pos + len(placeholder_id):]

        # 移除公式编号（如果存在）
        if equation_number and text_after:
            pattern = re.escape(equation_number)
            text_after = re.sub(r'[\s\t]*' + pattern + r'[\s\t]*', '', text_after)
            text_after = re.sub(r'\t+', ' ', text_after).strip()

        # 保存原始 run 的格式
        original_formatting = preserve_run_formatting(run)

        # 策略：拆分这个 run 为 3 部分
        # 1. text_before (保留原格式)
        # 2. 公式 (插入为独立段落或表格)
        # 3. text_after (保留原格式)

        # 获取段落的父元素
        para_element = para._element
        parent = para_element.getparent()
        para_index = parent.index(para_element)

        # 修改当前 run 为只包含 text_before
        run.text = text_before

        # 如果 text_before 为空，需要决定是否保留这个 run
        # 为了保持 run 结构，我们保留它但设为空文本

        # 在段落后插入公式表格
        if equation_number:
            # 创建公式表格（3列：空白、公式、编号）
            tbl = create_formula_table(math_run_element, equation_number)
            parent.insert(para_index + 1, tbl)
        else:
            # 创建只包含公式的简单段落
            formula_para = OxmlElement('w:p')
            # 复制段落属性
            pPr = para_element.find(qn('w:pPr'))
            if pPr is not None:
                from copy import deepcopy
                formula_para.append(deepcopy(pPr))
            # 添加公式 run
            formula_para.append(math_run_element)
            parent.insert(para_index + 1, formula_para)

        # 如果有 text_after，在当前段落中添加新 run
        if text_after:
            # 在当前 run 之后插入新 run
            new_run = para.add_run(text_after)
            apply_run_formatting(new_run, original_formatting)

            # 调整 run 顺序：新 run 应该紧跟在当前 run 之后
            # 获取所有 runs 的底层元素
            runs_elements = para_element.findall(qn('w:r'))
            if len(runs_elements) > run_idx + 1:
                # 移动最后一个 run（新添加的）到正确位置
                last_run_elem = runs_elements[-1]
                para_element.remove(last_run_elem)
                para_element.insert(run_idx + 1, last_run_elem)

        print(f"    ✓ Preserved {len(para.runs)} runs in paragraph")
        return True

    except Exception as e:
        print(f"    ✗ Error in replace_placeholder_in_run: {e}")
        import traceback
        traceback.print_exc()
        return False


def create_formula_table(math_run_element, equation_number):
    """
    创建公式表格（居中公式 + 右对齐编号）

    Args:
        math_run_element: OMML math run 元素
        equation_number: 公式编号

    Returns:
        OxmlElement: 表格元素
    """
    tbl = OxmlElement('w:tbl')

    # 表格属性 - 无边框，100% 宽度
    tblPr = OxmlElement('w:tblPr')

    # 表格宽度
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')  # 100%
    tblPr.append(tblW)

    # 无边框
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    tbl.append(tblPr)

    # 表格列
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(3):
        gridCol = OxmlElement('w:gridCol')
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # 表格行
    tr = OxmlElement('w:tr')

    # 单元格 1: 空（平衡）
    tc1 = OxmlElement('w:tc')
    tcPr1 = OxmlElement('w:tcPr')
    tcW1 = OxmlElement('w:tcW')
    tcW1.set(qn('w:type'), 'pct')
    tcW1.set(qn('w:w'), '250')  # 5%
    tcPr1.append(tcW1)
    tc1.append(tcPr1)
    p1 = OxmlElement('w:p')
    tc1.append(p1)
    tr.append(tc1)

    # 单元格 2: 公式（居中）
    tc2 = OxmlElement('w:tc')
    tcPr2 = OxmlElement('w:tcPr')
    tcW2 = OxmlElement('w:tcW')
    tcW2.set(qn('w:type'), 'pct')
    tcW2.set(qn('w:w'), '4000')  # 80%
    tcPr2.append(tcW2)
    tc2.append(tcPr2)

    p2 = OxmlElement('w:p')
    pPr2 = OxmlElement('w:pPr')
    jc2 = OxmlElement('w:jc')
    jc2.set(qn('w:val'), 'center')
    pPr2.append(jc2)
    p2.append(pPr2)
    p2.append(math_run_element)
    tc2.append(p2)
    tr.append(tc2)

    # 单元格 3: 编号（右对齐）
    tc3 = OxmlElement('w:tc')
    tcPr3 = OxmlElement('w:tcPr')
    tcW3 = OxmlElement('w:tcW')
    tcW3.set(qn('w:type'), 'pct')
    tcW3.set(qn('w:w'), '750')  # 15%
    tcPr3.append(tcW3)
    tc3.append(tcPr3)

    p3 = OxmlElement('w:p')
    if equation_number:
        pPr3 = OxmlElement('w:pPr')
        jc3 = OxmlElement('w:jc')
        jc3.set(qn('w:val'), 'right')
        pPr3.append(jc3)
        p3.append(pPr3)

        r3 = OxmlElement('w:r')
        t3 = OxmlElement('w:t')
        t3.text = equation_number
        r3.append(t3)
        p3.append(r3)

    tc3.append(p3)
    tr.append(tc3)

    tbl.append(tr)
    return tbl


def replace_formula_placeholders_preserve_runs(docx_path, formula_data,
                                                create_omml_func):
    """
    替换公式占位符，保留原始 run 结构

    这是对 hybrid_converter.py 中 _replace_formula_placeholders 的改进版本

    Args:
        docx_path: DOCX 文件路径
        formula_data: 公式数据字典 {page_num: [{'bbox': [...], 'latex': str, 'mathml': str}]}
        create_omml_func: 创建 OMML 的函数（从 MathML）

    Returns:
        int: 成功替换的公式数量
    """
    doc = Document(docx_path)

    # 创建占位符到公式数据的映射
    formula_map = {}
    for page_num, formulas in formula_data.items():
        for formula_idx, formula_info in enumerate(formulas):
            placeholder_id = f"FORMULA_PLACEHOLDER_{page_num}_{formula_idx}"
            formula_map[placeholder_id] = formula_info

    print(f"[Formula] Searching for {len(formula_map)} placeholders...")
    print(f"[Formula] PRESERVE MODE: Keeping original run structure")

    replaced_count = 0

    # 遍历所有段落
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text

        # 记录原始 run 数量
        original_run_count = len(para.runs)

        # 查找占位符
        for placeholder_id, formula_info in formula_map.items():
            if placeholder_id in para_text:
                mathml = formula_info['mathml']
                latex = formula_info['latex']
                equation_number = formula_info.get('equation_number', None)

                print(f"\n  [Formula] Found '{placeholder_id}' in para {para_idx}")
                print(f"    Original runs: {original_run_count}")

                # 查找包含占位符的具体 run
                run_idx, run, start_pos = find_run_with_placeholder(para, placeholder_id)

                if run is None:
                    print(f"    ✗ Could not locate placeholder in runs")
                    continue

                print(f"    Located in run {run_idx} at position {start_pos}")

                try:
                    # 去除 MathML 声明
                    if mathml.startswith('<?xml'):
                        mathml = mathml[mathml.index('?>') + 2:].strip()

                    # 创建 OMML 元素
                    math_run = create_omml_func(mathml)

                    if math_run:
                        # 保留 run 结构的替换
                        success = replace_placeholder_in_run(
                            para, run_idx, run, start_pos,
                            placeholder_id, math_run, equation_number
                        )

                        if success:
                            replaced_count += 1
                            new_run_count = len(para.runs)
                            print(f"    New runs: {new_run_count} (delta: {new_run_count - original_run_count})")
                            print(f"    ✓ Replaced: {latex[:60]}...")
                        else:
                            print(f"    ✗ Failed to replace")
                    else:
                        print(f"    ✗ Failed to create OMML")

                except Exception as e:
                    print(f"    ✗ Error: {e}")
                    import traceback
                    traceback.print_exc()

                break  # 移到下一个段落

    print(f"\n[Formula] Replaced {replaced_count}/{len(formula_map)} formulas")

    # 保存文档
    doc.save(docx_path)
    print(f"[Formula] Document saved with preserved run structure")

    return replaced_count


# 使用示例
if __name__ == "__main__":
    print("=" * 80)
    print("改进的公式插入方法 - 保留 run 结构")
    print("=" * 80)
    print()
    print("核心改进：")
    print("1. 不使用 para.clear() 破坏所有 runs")
    print("2. 精确定位包含占位符的 run")
    print("3. 只替换特定 run，保留其他 runs 的格式")
    print("4. 保持段落的原始 run 数量和格式属性")
    print()
    print("预期效果：")
    print("  - pdf2docx 创建的丰富 run 结构得以保留")
    print("  - 多种字体、样式的 runs 不会被合并")
    print("  - 段落 run 数量接近 smallpdf (30-113 runs/para)")
    print("  - 而不是当前的 3-11 runs/para")
    print()
    print("=" * 80)
