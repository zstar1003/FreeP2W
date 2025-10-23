"""
Integrate formatting into hybrid conversion workflow
"""

from hybrid_converter import HybridConverter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
import re
import os

def format_converted_docx(docx_path):
    """
    Apply formatting to converted DOCX:
    1. Unify font sizes to 10pt
    2. Set alignments: sections=LEFT, figures=CENTER, body=JUSTIFY
    3. Clean TOC spacing
    """

    doc = Document(docx_path)

    print("\n[Format] Applying final formatting...")

    # Determine body text size
    font_sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                font_sizes.append(run.font.size.pt)

    body_size = Counter(font_sizes).most_common(1)[0][0] if font_sizes else 10.0
    print(f"[Format] Body text size: {body_size}pt")

    # Detect TOC
    toc_start = None
    toc_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        if toc_start is None and ('contents' in text or 'table of contents' in text):
            toc_start = i

        if toc_start is not None and toc_end is None and i > toc_start:
            para_text = para.text.strip()
            if (re.match(r'^\d+\.?\d*\s+[A-Z]', para_text) or
                any(keyword in text for keyword in ['abstract', 'introduction']) or
                len(para_text) > 100):
                toc_end = i
                break

    if toc_start and not toc_end:
        toc_end = min(toc_start + 15, len(doc.paragraphs))

    # Format paragraphs
    stats = {'titles': 0, 'sections': 0, 'figures': 0, 'body': 0, 'toc': 0}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Detect paragraph type
        max_font_size = max((run.font.size.pt for run in para.runs if run.font.size), default=0)

        is_figure = text.lower().startswith(('figure', 'fig.', 'fig '))
        is_section = bool(re.match(r'^[\d\.]+\s+[A-Z]', text))
        is_title = max_font_size > 12 or (len(text) < 100 and text.isupper() and not is_figure)
        in_toc = toc_start is not None and toc_start <= i < toc_end

        # Apply formatting
        if in_toc:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Clean spaces
            cleaned = re.sub(r'\s+', ' ', text).strip()
            if cleaned != text and para.runs:
                para.clear()
                run = para.add_run(cleaned)
                run.font.size = Pt(body_size)
            stats['toc'] += 1

        elif is_figure:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                if run.font.size and run.font.size.pt != body_size:
                    run.font.size = Pt(body_size)
            stats['figures'] += 1

        elif is_section:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats['sections'] += 1

        elif is_title:
            if not para.alignment or para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats['titles'] += 1

        else:
            # Body text
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                if run.font.size and run.font.size.pt <= 12 and run.font.size.pt != body_size:
                    run.font.size = Pt(body_size)
            stats['body'] += 1

    print(f"[Format] Formatted {stats['sections']} sections, {stats['figures']} figures, {stats['body']} body paragraphs")

    # Save
    doc.save(docx_path)
    print(f"[Format] Formatting completed")


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python format_and_convert.py <pdf_file> [output.docx]")
        sys.exit(1)

    pdf_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else pdf_file.replace('.pdf', '_converted.docx')

    print("=" * 80)
    print("PDF TO DOCX CONVERSION WITH FORMATTING")
    print("=" * 80)

    # Step 1: Convert PDF to DOCX
    print("\nStep 1: Converting PDF to DOCX...")
    converter = HybridConverter()
    converter.convert(pdf_file, output_file, use_yolo_images=True)

    # Step 2: Apply formatting
    print("\nStep 2: Applying formatting...")
    format_converted_docx(output_file)

    print("\n" + "=" * 80)
    print(f"COMPLETED!")
    print(f"Output: {output_file}")
    print("=" * 80)
