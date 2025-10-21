"""
Hybrid PDF to DOCX Converter using DocLayout-YOLO + pdf2docx

Strategy:
1. Use DocLayout-YOLO to detect layout regions (especially Figure)
2. Extract Figure regions as complete images
3. Use pdf2docx for text content
4. Assemble into DOCX with proper ordering
"""

import os
import torch
import fitz
from docx import Document
from docx.shared import Inches, Pt
from doclayout_yolo import YOLOv10
import io


class DocLayoutYOLOConverter:
    """PDF to DOCX converter using DocLayout-YOLO for layout detection"""

    def __init__(self, model_path="weights/doclayout_yolo_docstructbench_imgsz1024.pt"):
        """
        Initialize converter with DocLayout-YOLO model

        Args:
            model_path: Path to DocLayout-YOLO model weights
        """
        print(f"[INFO] Loading DocLayout-YOLO model...")

        # Automatically select device
        self.device = 'cuda' if torch.cuda.is_available() else 'mps' if torch.backends.mps.is_available() else 'cpu'
        print(f"[INFO] Using device: {self.device}")

        # Load model
        self.model = YOLOv10(model_path)
        print(f"[INFO] Model loaded successfully")

    def convert(self, pdf_path, docx_path, imgsz=1024, conf=0.25, dpi=150, image_quality=4.0):
        """
        Convert PDF to DOCX using layout detection

        Args:
            pdf_path: Path to input PDF
            docx_path: Path to output DOCX
            imgsz: Image size for YOLO detection (default: 1024)
            conf: Confidence threshold (default: 0.25)
            dpi: DPI for PDF rendering (default: 150)
            image_quality: Quality multiplier for extracted images (default: 4.0)
        """
        print(f"\n{'='*80}")
        print(f"Converting: {pdf_path}")
        print(f"{'='*80}\n")

        # Open PDF
        pdf_doc = fitz.open(pdf_path)
        docx_doc = Document()

        # Process each page
        for page_num in range(len(pdf_doc)):
            print(f"\n{'-'*80}")
            print(f"Processing Page {page_num + 1}/{len(pdf_doc)}")
            print(f"{'-'*80}")

            page = pdf_doc[page_num]

            # Step 1: Detect layout with DocLayout-YOLO
            layout_regions = self._detect_layout(page, imgsz, conf, dpi)

            # Step 2: Sort regions by reading order (top to bottom, left to right)
            layout_regions.sort(key=lambda r: (r['bbox'][1], r['bbox'][0]))

            # Step 3: Process each region
            for i, region in enumerate(layout_regions, 1):
                region_type = region['type']
                bbox = region['bbox']
                confidence = region['confidence']

                print(f"  [{i}/{len(layout_regions)}] {region_type} (conf: {confidence:.2f})")

                if region_type in ['figure', 'Figure', 'picture', 'Picture']:
                    # Extract figure as image
                    self._insert_figure_region(page, bbox, docx_doc, image_quality)

                elif region_type in ['table', 'Table']:
                    # Extract table as image (for now)
                    print(f"      [Table] Extracting as image")
                    self._insert_figure_region(page, bbox, docx_doc, image_quality)

                elif region_type in ['text', 'Text', 'title', 'Title', 'caption', 'Caption']:
                    # Extract text
                    self._insert_text_region(page, bbox, docx_doc, region_type)

                else:
                    print(f"      [Skipped] Region type: {region_type}")

            # Add page break
            if page_num < len(pdf_doc) - 1:
                docx_doc.add_page_break()

        # Save DOCX
        pdf_doc.close()
        docx_doc.save(docx_path)

        print(f"\n{'='*80}")
        print(f"Conversion completed!")
        print(f"Output: {docx_path}")
        print(f"{'='*80}\n")

    def _detect_layout(self, page, imgsz, conf, dpi):
        """
        Detect layout regions using DocLayout-YOLO

        Args:
            page: fitz.Page object
            imgsz: Image size for YOLO
            conf: Confidence threshold
            dpi: DPI for rendering

        Returns:
            list: Detected regions with type, bbox, confidence
        """
        # Convert PDF page to image
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")

        # Save to temporary file for YOLO
        temp_img_path = "_temp_page.png"
        with open(temp_img_path, "wb") as f:
            f.write(img_bytes)

        # Run YOLO detection
        det_res = self.model.predict(
            temp_img_path,
            imgsz=imgsz,
            conf=conf,
            device=self.device,
        )

        # Clean up temp file
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)

        # Parse results
        regions = []

        if len(det_res) > 0 and hasattr(det_res[0], 'boxes'):
            boxes = det_res[0].boxes

            for box in boxes:
                # Get bbox coordinates (in image pixels)
                x1, y1, x2, y2 = box.xyxy[0].cpu().numpy()

                # Get class and confidence
                cls_id = int(box.cls[0].cpu().numpy())
                conf_score = float(box.conf[0].cpu().numpy())

                # Get class name
                class_name = self.model.names[cls_id]

                # Convert from image coordinates to PDF coordinates
                scale_x = page.rect.width / pix.width
                scale_y = page.rect.height / pix.height

                pdf_bbox = [
                    x1 * scale_x,
                    y1 * scale_y,
                    x2 * scale_x,
                    y2 * scale_y
                ]

                regions.append({
                    'type': class_name,
                    'bbox': pdf_bbox,
                    'confidence': conf_score
                })

        print(f"  Detected {len(regions)} regions")

        # Show summary by type
        type_counts = {}
        for r in regions:
            type_counts[r['type']] = type_counts.get(r['type'], 0) + 1

        for rtype, count in sorted(type_counts.items()):
            print(f"    - {rtype}: {count}")

        return regions

    def _insert_figure_region(self, page, bbox, docx_doc, quality=4.0):
        """
        Extract figure region from PDF and insert into DOCX

        Args:
            page: fitz.Page object
            bbox: Bounding box [x0, y0, x1, y1]
            docx_doc: Document object
            quality: Image quality multiplier
        """
        # Create rect from bbox
        rect = fitz.Rect(bbox)

        # Clip page region as high-quality image
        mat = fitz.Matrix(quality, quality)
        pix = page.get_pixmap(matrix=mat, clip=rect)

        # Convert to PNG bytes
        img_data = pix.tobytes("png")

        # Insert into DOCX
        try:
            img_stream = io.BytesIO(img_data)

            # Calculate width in inches
            max_width = 6.0  # inches
            width_pt = rect.width
            width_inches = width_pt / 72.0

            if width_inches > max_width:
                width_inches = max_width

            docx_doc.add_picture(img_stream, width=Inches(width_inches))

            print(f"      [Figure] Inserted: {rect.width:.1f} x {rect.height:.1f} pt")

        except Exception as e:
            print(f"      [Error] Failed to insert figure: {e}")

    def _insert_text_region(self, page, bbox, docx_doc, region_type):
        """
        Extract text region from PDF and insert into DOCX

        Args:
            page: fitz.Page object
            bbox: Bounding box [x0, y0, x1, y1]
            docx_doc: Document object
            region_type: Type of text region
        """
        # Extract text from region
        rect = fitz.Rect(bbox)
        text = page.get_text("text", clip=rect).strip()

        if not text:
            return

        # Add to DOCX with appropriate styling
        if region_type.lower() in ['title', 'heading']:
            para = docx_doc.add_heading(text, level=1)
        elif region_type.lower() in ['caption']:
            para = docx_doc.add_paragraph(text)
            para.runs[0].font.size = Pt(9)
            para.runs[0].italic = True
        else:
            para = docx_doc.add_paragraph(text)

        print(f"      [Text] Inserted: {len(text)} chars")


def main():
    """Test the converter"""
    import argparse

    parser = argparse.ArgumentParser(description='PDF to DOCX using DocLayout-YOLO')
    parser.add_argument('--pdf', required=True, help='Input PDF file')
    parser.add_argument('--output', help='Output DOCX file (default: input_doclayout.docx)')
    parser.add_argument('--model', default='weights/doclayout_yolo_docstructbench_imgsz1024.pt', help='Model path')
    parser.add_argument('--imgsz', type=int, default=1024, help='YOLO image size')
    parser.add_argument('--conf', type=float, default=0.25, help='Confidence threshold')
    parser.add_argument('--dpi', type=int, default=150, help='DPI for PDF rendering')
    parser.add_argument('--image-quality', type=float, default=4.0, help='Image quality multiplier')

    args = parser.parse_args()

    # Set default output path
    if args.output is None:
        args.output = args.pdf.replace('.pdf', '_doclayout.docx')

    # Create converter
    converter = DocLayoutYOLOConverter(model_path=args.model)

    # Convert
    converter.convert(
        pdf_path=args.pdf,
        docx_path=args.output,
        imgsz=args.imgsz,
        conf=args.conf,
        dpi=args.dpi,
        image_quality=args.image_quality
    )


if __name__ == "__main__":
    main()
