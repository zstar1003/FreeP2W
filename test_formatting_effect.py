"""
快速测试格式化效果 - 对比分析

对比 result.docx 和 result_formatted.docx 的格式差异
"""

from docx import Document
from collections import Counter


def analyze_formatting(docx_path, doc_name):
    """分析文档格式"""
    doc = Document(docx_path)

    stats = {
        'total_paragraphs': len(doc.paragraphs),
        'fonts': [],
        'font_sizes': [],
        'alignments': [],
        'body_paragraphs': 0
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 20:
            continue

        # 记录对齐方式
        if para.alignment is not None:
            stats['alignments'].append(str(para.alignment))
        else:
            stats['alignments'].append('None')

        # 判断是否为正文（简单判断）
        is_title = False
        for run in para.runs:
            if run.font.size and run.font.size.pt >= 14:
                is_title = True
                break
            if run.font.bold and len(text) < 100:
                is_title = True
                break

        if not is_title and 'Figure' not in text and 'Table' not in text:
            stats['body_paragraphs'] += 1

            # 记录正文的字体和大小
            for run in para.runs:
                if run.font.name:
                    stats['fonts'].append(run.font.name)
                else:
                    stats['fonts'].append('None')

                if run.font.size:
                    stats['font_sizes'].append(run.font.size.pt)

    print(f"\n{'='*80}")
    print(f"{doc_name}")
    print(f"{'='*80}")
    print(f"总段落数: {stats['total_paragraphs']}")
    print(f"正文段落数（估算）: {stats['body_paragraphs']}")

    if stats['fonts']:
        font_counter = Counter(stats['fonts'])
        print(f"\n正文字体分布:")
        for font, count in font_counter.most_common(5):
            pct = count / len(stats['fonts']) * 100
            print(f"  {font}: {count} ({pct:.1f}%)")

    if stats['font_sizes']:
        size_counter = Counter(stats['font_sizes'])
        print(f"\n正文字体大小分布:")
        for size, count in size_counter.most_common(5):
            pct = count / len(stats['font_sizes']) * 100
            print(f"  {size}pt: {count} ({pct:.1f}%)")

    align_counter = Counter(stats['alignments'])
    print(f"\n段落对齐方式分布:")
    for align, count in align_counter.most_common():
        pct = count / stats['total_paragraphs'] * 100
        align_name = align.split('.')[-1] if '.' in align else align
        print(f"  {align_name}: {count} ({pct:.1f}%)")

    return stats


if __name__ == "__main__":
    import os

    print("="*80)
    print("格式化效果对比分析")
    print("="*80)

    # 检查文件是否存在
    if not os.path.exists('result.docx'):
        print("\n错误: result.docx 不存在")
        print("请先运行 PDF 转换生成 result.docx")
        exit(1)

    if not os.path.exists('result_formatted.docx'):
        print("\n警告: result_formatted.docx 不存在")
        print("将只分析 result.docx")
        print("\n运行以下命令生成格式化文件:")
        print("  python post_format_english_body.py")
        print()

        stats1 = analyze_formatting('result.docx', 'result.docx (原始)')
        exit(0)

    # 对比分析
    stats1 = analyze_formatting('result.docx', 'result.docx (原始)')
    stats2 = analyze_formatting('result_formatted.docx', 'result_formatted.docx (格式化后)')

    # 对比总结
    print(f"\n{'='*80}")
    print("对比总结")
    print(f"{'='*80}")

    # 字体统一度
    fonts1 = [f for f in stats1['fonts'] if f != 'None']
    fonts2 = [f for f in stats2['fonts'] if f != 'None']

    if fonts1:
        unique_fonts1 = len(set(fonts1))
        print(f"\n原始文档正文字体种类: {unique_fonts1}")
    else:
        print(f"\n原始文档正文字体: 大多未设置")

    if fonts2:
        unique_fonts2 = len(set(fonts2))
        times_new_roman_count = fonts2.count('Times New Roman')
        times_pct = times_new_roman_count / len(fonts2) * 100 if fonts2 else 0
        print(f"格式化后正文字体种类: {unique_fonts2}")
        print(f"Times New Roman 占比: {times_pct:.1f}%")

    # 字体大小统一度
    if stats1['font_sizes']:
        sizes1 = Counter(stats1['font_sizes'])
        most_common_size1 = sizes1.most_common(1)[0]
        print(f"\n原始文档最常见字体大小: {most_common_size1[0]}pt ({most_common_size1[1]} 次)")

    if stats2['font_sizes']:
        sizes2 = Counter(stats2['font_sizes'])
        most_common_size2 = sizes2.most_common(1)[0]
        size_10_count = sizes2.get(10.0, 0)
        size_10_pct = size_10_count / sum(sizes2.values()) * 100 if sizes2 else 0
        print(f"格式化后最常见字体大小: {most_common_size2[0]}pt ({most_common_size2[1]} 次)")
        print(f"10pt 占比: {size_10_pct:.1f}%")

    # 对齐方式
    align1 = Counter(stats1['alignments'])
    align2 = Counter(stats2['alignments'])

    justify1 = sum(count for align, count in align1.items() if 'JUSTIFY' in align)
    justify2 = sum(count for align, count in align2.items() if 'JUSTIFY' in align)

    justify_pct1 = justify1 / stats1['total_paragraphs'] * 100
    justify_pct2 = justify2 / stats2['total_paragraphs'] * 100

    print(f"\n原始文档两端对齐段落: {justify1} ({justify_pct1:.1f}%)")
    print(f"格式化后两端对齐段落: {justify2} ({justify_pct2:.1f}%)")

    print(f"\n{'='*80}")
    print("结论:")
    print(f"{'='*80}")

    if times_pct > 80:
        print("[优秀] 字体统一效果: Times New Roman 占比 > 80%")
    elif times_pct > 60:
        print("[良好] 字体统一效果: Times New Roman 占比 > 60%")
    else:
        print("[待改进] 字体统一效果: Times New Roman 占比 < 60%")

    if size_10_pct > 80:
        print("[优秀] 字体大小统一效果: 10pt 占比 > 80%")
    elif size_10_pct > 60:
        print("[良好] 字体大小统一效果: 10pt 占比 > 60%")
    else:
        print("[待改进] 字体大小统一效果: 10pt 占比 < 60%")

    if justify_pct2 - justify_pct1 > 20:
        print("[优秀] 对齐统一效果: 两端对齐显著增加")
    elif justify_pct2 > justify_pct1:
        print("[良好] 对齐统一效果: 两端对齐有所增加")
    else:
        print("[保持] 对齐统一效果: 保持原样")

    print(f"{'='*80}")
